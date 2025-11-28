#!/usr/bin/env python3
"""
Generation of Section 2 "SYSTEM ARCHITECTURE" in DOCX format with academic style (v2).
This version minimizes bullet lists and uses prose paragraphs for better academic readability.
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess

# Base paths
BASE_DIR = Path(__file__).parent
DIAGRAMS_DIR = BASE_DIR / "all_diagrams" / "updated_without_title_eng"
OUTPUT_FILE = BASE_DIR / "Section2_Architecture_EN_v2.docx"

# Diagrams for insertion (English versions without title)
DIAGRAMS = {
    "fig2": DIAGRAMS_DIR / "01_system_architecture.svg",
    "fig2a": DIAGRAMS_DIR / "07_schema_validation_architecture.svg",
    "fig2b": DIAGRAMS_DIR / "08_query_syntax_analyzer.svg",
    "fig2c": DIAGRAMS_DIR / "09_query_antipattern_analyzer.svg",
    "fig3": DIAGRAMS_DIR / "02_data_flow.svg",
    "fig4": DIAGRAMS_DIR / "04_protocol_based_architecture.svg",
    "fig5a": DIAGRAMS_DIR / "05_db_identity_normalization.svg",
    "fig6": DIAGRAMS_DIR / "03_llm_judge_architecture.svg",
    "fig6a": DIAGRAMS_DIR / "06_llm_prompt_structure.svg",
}


def convert_svg_to_png(svg_path: Path, png_path: Path, dpi: int = 150):
    """Convert SVG to PNG using Inkscape."""
    try:
        cmd = [
            "inkscape",
            str(svg_path),
            "--export-type=png",
            f"--export-dpi={dpi}",
            f"--export-filename={png_path}",
        ]
        subprocess.run(cmd, check=True, capture_output=True)
        return True
    except (subprocess.CalledProcessError, FileNotFoundError) as e:
        print(f"Conversion error for {svg_path.name}: {e}")
        return False


def add_heading(doc, text, level=1):
    """Add a heading with formatting."""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.name = "Times New Roman"
    heading.runs[0].font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    heading.runs[0].font.bold = True
    return heading


def add_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add a paragraph with formatting."""
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
    return para


def add_code_block(doc, code_text):
    """Add a code block in monospaced font."""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    return para


def add_image_from_svg(doc, svg_path: Path, caption: str, width: float = 6.0):
    """Convert SVG to PNG and insert image with caption."""
    png_path = svg_path.with_suffix(".png")

    # Convert SVG to PNG if needed
    if not png_path.exists() or png_path.stat().st_mtime < svg_path.stat().st_mtime:
        print(f"Converting {svg_path.name}...")
        if not convert_svg_to_png(svg_path, png_path):
            print(f"  ⚠️ Failed to convert {svg_path.name}, skipping")
            add_paragraph(doc, f"[Image placeholder: {caption}]", italic=True)
            return

    # Insert image
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(png_path), width=Inches(width))

        # Add caption
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].font.size = Pt(10)
        caption_para.runs[0].font.italic = True
        print(f"  ✅ Added {svg_path.name}")
    except Exception as e:
        print(f"  ⚠️ Error inserting image {svg_path.name}: {e}")
        add_paragraph(doc, f"[Image placeholder: {caption}]", italic=True)


def generate_section2_docx():
    """Generate DOCX document with Section 2 in academic style."""
    print("🚀 Generating Section 2 'SYSTEM ARCHITECTURE' (Academic v2)...")

    # Create document
    doc = Document()

    # Configure styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Section title
    title = doc.add_heading("2. SYSTEM ARCHITECTURE", level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True

    # 2.1. High-level overview
    add_heading(doc, "2.1. Architectural Overview and Design Principles", level=2)
    
    add_paragraph(
        doc,
        "The text2sql-pipeline system represents a streaming data-processing architecture designed to address "
        "the challenges of large-scale Text-to-SQL dataset validation and quality assessment. The architecture "
        "implements three foundational design principles: modularity through protocol-based component interfaces, "
        "scalability via constant-memory streaming processing, and extensibility through automatic dependency injection. "
        "This design enables the system to process datasets of arbitrary size while maintaining deterministic memory "
        "footprint characteristics, a critical requirement when analyzing multi-gigabyte benchmark datasets such as Spider, "
        "BIRD, or WikiSQL.",
    )

    add_paragraph(
        doc,
        "The implementation leverages a carefully selected technology stack centered around Python 3, chosen for its "
        "rich ecosystem of SQL processing libraries and extensive support for language model integration. DuckDB serves "
        "as the analytics backend, providing columnar storage optimized for analytical queries while maintaining compatibility "
        "with standard SQL interfaces. For SQL parsing and analysis, the system employs sqlglot, a comprehensive dialect-aware "
        "parser capable of handling the syntactic variations across SQLite, PostgreSQL, MySQL, and other SQL variants commonly "
        "found in Text-to-SQL datasets. Component orchestration relies on the dependency_injector library, which facilitates "
        "type-safe dependency resolution and enables runtime composition of analysis pipelines through declarative configuration.",
    )

    add_paragraph(
        doc,
        "User interaction with the system occurs primarily through a command-line interface exposed via the `text2sql` "
        "entry point. The interface supports two distinct operational modes corresponding to the system's primary use cases. "
        "The first mode, invoked through `text2sql run --config <config.yaml>`, executes the complete analysis pipeline from "
        "raw dataset ingestion through all configured analyzers to final metric persistence and report generation. The second "
        "mode, accessible via `text2sql report`, enables standalone report generation from previously computed metrics, "
        "supporting both configuration-driven batch report generation (`--config`) and targeted single-report generation "
        "(`--database <path> --type <report-type>`). This separation of concerns between analysis and reporting allows "
        "researchers to iterate on report designs and formats without re-executing computationally expensive analysis stages, "
        "particularly relevant when semantic validation via large language models can require hours of processing time.",
    )

    add_paragraph(
        doc,
        "The architectural design embodies several key principles that distinguish it from traditional batch-processing "
        "approaches commonly employed in dataset analysis tools. Streaming processing ensures that records flow through "
        "the pipeline individually, with each record being fully processed before the next is loaded, thereby maintaining "
        "O(1) memory complexity regardless of dataset cardinality. The protocol-based architecture establishes clear "
        "contracts between components through Python's Protocol typing construct, enabling loose coupling between "
        "implementation and interface while preserving type safety guarantees. Dependency injection automates the wiring "
        "of complex component graphs, reducing boilerplate configuration code and facilitating unit testing through "
        "straightforward mock object substitution. Finally, multi-dialect database support acknowledges the heterogeneous "
        "nature of Text-to-SQL datasets, where different benchmarks may utilize different SQL dialects, necessitating "
        "dialect-aware parsing and validation strategies.",
    )

    # 2.2. Five-layer architecture
    add_heading(doc, "2.2. Five-Layer System Architecture", level=2)
    
    add_paragraph(
        doc,
        "The system architecture follows a layered design pattern comprising five distinct layers, each responsible for "
        "a specific aspect of the data processing pipeline (Figure 2). This layered organization promotes separation of "
        "concerns, facilitates independent testing and development of components, and enables selective activation or "
        "deactivation of functionality based on analysis requirements.",
    )

    # Figure 2
    add_image_from_svg(
        doc,
        DIAGRAMS["fig2"],
        "Figure 2. Five-layer architecture of the text2sql-pipeline system",
        width=6.5,
    )

    add_paragraph(
        doc,
        "The Input Layer constitutes the foundation of the pipeline, responsible for abstracting diverse data sources "
        "behind a uniform iterator interface. The layer implements four specialized loaders, each optimized for a "
        "particular data format commonly encountered in Text-to-SQL research. The JSONL Loader provides streaming access "
        "to line-delimited JSON files, a format widely adopted for its balance between human readability and machine "
        "efficiency, with additional support for gzip compression to reduce storage overhead. The CSV Loader incorporates "
        "automatic delimiter detection and encoding inference, handling the format variations frequently encountered in "
        "dataset downloads and conversions. Standard JSON array processing is handled by the JSON Loader, which accommodates "
        "datasets distributed in traditional JSON format despite the memory implications of loading complete arrays. "
        "Finally, the HuggingFace Loader integrates directly with the HuggingFace Datasets Hub, enabling immediate access "
        "to the growing collection of publicly available Text-to-SQL datasets without requiring manual download and "
        "format conversion steps. Critically, all loaders implement Python's iterator protocol, yielding individual records "
        "on demand rather than materializing entire datasets in memory, thereby ensuring constant memory consumption "
        "regardless of dataset magnitude.",
    )

    add_paragraph(
        doc,
        "The Normalization Layer performs essential data standardization operations required to reconcile the diverse "
        "structural and semantic conventions employed across different Text-to-SQL datasets. This layer operates through "
        "three sequential normalizers, each addressing a specific aspect of data heterogeneity. The Alias Mapper component "
        "resolves field name variations, mapping semantically equivalent fields with different names (such as 'query' versus "
        "'sql', or 'db_id' versus 'dbId') to a canonical internal representation, thereby eliminating dataset-specific field "
        "naming concerns from downstream components. The ID Assignment normalizer ensures that every record possesses a unique "
        "identifier, generating deterministic IDs based on either incremental counters or cryptographic hashing of record content, "
        "facilitating cross-reference between the annotated output dataset and the metrics database. The DB Identity Normalizer "
        "implements sophisticated logic to guarantee that every record is associated with a valid database identifier, either "
        "by verifying the health of an explicitly provided database ID or by automatically creating a database from an embedded "
        "DDL schema specification, then computing a deterministic ID from that schema through cryptographic hashing. The output "
        "of this layer is a standardized DataItem object with guaranteed presence of essential fields: id, question, sql, dbId, "
        "an optional schema field, and a metadata dictionary for auxiliary information.",
    )

    add_paragraph(
        doc,
        "The Analysis Layer represents the computational core of the system, housing five independent analyzers that examine "
        "different quality dimensions of Text-to-SQL data. The Schema Validation Analyzer scrutinizes database schema integrity, "
        "detecting four categories of foreign key errors (missing referenced tables, missing referenced columns, arity mismatches "
        "between referencing and referenced columns, and foreign keys targeting non-key columns), identifying case-insensitive "
        "column name duplicates within tables, validating data type compatibility with the target SQL dialect, and implementing "
        "a database-ID-keyed caching strategy to amortize the cost of schema introspection across multiple queries referencing "
        "the same database. The Query Syntax Analyzer performs static structural analysis of SQL statements, constructing abstract "
        "syntax trees through the sqlglot parser, computing a normalized complexity metric on a 0-100 scale based on the presence "
        "and nesting depth of various SQL constructs, extracting structural features including JOIN counts, subquery depths, "
        "common table expression usage, aggregation functions, and window functions, and classifying queries into difficulty "
        "categories (Easy, Medium, Hard, Expert) to facilitate analysis of dataset complexity distribution.",
    )

    add_paragraph(
        doc,
        "The Query Execution Analyzer evaluates dynamic executability of SQL statements within their associated database contexts, "
        "executing queries in a sandboxed SELECT-only mode to prevent data modification, measuring execution time with microsecond "
        "precision to identify performance outliers, counting returned row counts to detect empty result sets that may indicate "
        "semantic errors, and classifying execution failures into categories to distinguish syntax errors from semantic errors and "
        "database connection issues. The Query Antipattern Detector implements a rule-based system for identifying 14 common SQL "
        "antipatterns including unbounded SELECT statements lacking LIMIT clauses, SELECT * projections that reduce query "
        "maintainability, functions applied in WHERE clauses that prevent index utilization, correlated subqueries with "
        "quadratic complexity characteristics, and implicit joins lacking explicit JOIN syntax, with each detected pattern "
        "assigned a severity weight contributing to an overall quality score ranging from 0 to 100.",
    )

    add_paragraph(
        doc,
        "The Semantic LLM Judge Analyzer represents the most sophisticated component, employing large language models to assess "
        "semantic correspondence between natural language questions and their associated SQL queries. This analyzer implements a "
        "multi-model consensus voting architecture supporting concurrent querying of OpenAI GPT models, Anthropic Claude models, "
        "Google Gemini models, and locally-hosted Ollama models. The system generates minimal DDL schemas by parsing SQL queries "
        "to extract referenced tables, querying database catalogs for metadata on only those referenced tables, sampling a "
        "configurable number of example values from each column (default: 2 examples per column), and formatting CREATE TABLE "
        "statements with inline comments containing these examples, thereby providing language models with concrete data context "
        "while reducing token consumption by factors of 5-10 compared to transmitting full database schemas. Model responses, "
        "structured as verdict classifications (CORRECT, PARTIALLY_CORRECT, INCORRECT, UNANSWERABLE) with confidence scores and "
        "detailed reasoning, are aggregated through weighted voting to determine consensus classifications and identify cases "
        "where models disagree, potentially indicating ambiguous or problematic examples.",
    )

    add_paragraph(
        doc,
        "The Output Layer manages result persistence through two parallel write streams, reflecting the dual objectives of dataset "
        "annotation and aggregate analytics. The Annotated Dataset stream augments original dataset records with metadata from all "
        "executed analyzers, preserving the complete analysis provenance in JSONL format for subsequent inspection, filtering, or "
        "machine learning preprocessing. The DuckDB Analytics Database stream accumulates structured metrics from all analyzers in a "
        "columnar OLAP database, enabling complex analytical queries across the full dataset using standard SQL syntax, supporting "
        "export to Pandas or Polars dataframes for statistical analysis or visualization, and facilitating Parquet export for "
        "long-term archival or integration with big data processing frameworks. This dual-output strategy ensures that both "
        "record-level details and aggregate statistics remain accessible without requiring re-execution of the analysis pipeline.",
    )

    add_paragraph(
        doc,
        "The Report Generation Layer provides an optional visualization tier that synthesizes DuckDB metrics into human-readable "
        "Markdown documents. This layer is explicitly designed as optional, configurable via the reports.enabled flag, to accommodate "
        "workflows where rapid iteration on large datasets necessitates skipping report generation during initial exploration phases. "
        "When enabled, the layer generates seven specialized reports covering summary statistics aggregated across all analyzers, "
        "detailed schema validation findings enumerating databases with integrity violations, semantic validation issues highlighting "
        "queries judged incorrect or unanswerable by language models, execution failures cataloging queries that could not be run "
        "successfully, structural complexity profiles showing the distribution of query difficulty levels and feature utilization, "
        "table coverage analysis identifying underutilized or overrepresented tables in the dataset, and quality assessments "
        "summarizing antipattern detection results. The architectural separation between metrics persistence (mandatory) and report "
        "generation (optional) offers several advantages: DuckDB serves as the canonical source of truth for all metrics with "
        "unrestricted SQL query access, reports can be regenerated at any time from the metrics database without re-analyzing the "
        "dataset, users can author custom reports through ad-hoc SQL queries against the metrics database, and the performance "
        "impact of report generation can be eliminated when analyzing large datasets under time constraints.",
    )

    # 2.3. Schema Validation Analyzer
    doc.add_page_break()
    add_heading(doc, "2.3. Schema Validation Analyzer", level=2)
    
    add_paragraph(
        doc,
        "The Schema Validation Analyzer constitutes a critical component of the quality assessment pipeline, scrutinizing "
        "database schema integrity to detect structural errors and data quality issues that could compromise query execution "
        "correctness or semantic validity. Unlike syntax analysis which examines SQL query correctness or execution testing "
        "which validates runtime behavior, schema validation focuses on the foundational layer—the database structure itself—"
        "identifying problems in table definitions, constraint specifications, and referential integrity that affect all queries "
        "operating against those databases. The analyzer implements a dual-level validation architecture distinguishing between "
        "structural errors in schema declarations (design flaws preventing proper constraint enforcement) and data violations "
        "in existing rows (integrity breaches where data contradicts declared constraints), enabling both proactive schema "
        "repair and targeted data cleaning strategies.",
    )

    add_image_from_svg(
        doc,
        DIAGRAMS["fig2a"],
        "Figure 2a. Schema validation architecture with dual-level validation approach",
        width=6.5,
    )

    add_paragraph(
        doc,
        "Figure 2a illustrates the three-branch schema validation architecture. The analyzer processes databases through parallel "
        "validation pathways, each targeting distinct quality dimensions. The first branch extracts foreign key definitions from "
        "database metadata through dialect-specific system catalogs, with fallback to dataset-provided schema files when explicit "
        "constraint metadata proves unavailable. The second branch performs schema introspection via column enumeration and primary "
        "key discovery to detect structural anomalies. The third branch executes data queries to assess content quality. This "
        "multi-pathway design enables comprehensive validation spanning constraint integrity, column definitions, and data completeness.",
    )

    add_paragraph(
        doc,
        "Foreign key structural validation examines constraint definitions through five checks: Missing Parent Table (verifying "
        "referenced tables exist), Missing Parent Column (confirming parent columns presence), Arity Mismatch (enforcing equal column "
        "counts in composite keys), Type Mismatch (detecting incompatible types after normalization to dialect affinity families), and "
        "Target Not Key (verifying parent columns possess uniqueness constraints). Foreign key data validation detects orphaned values—"
        "rows where foreign key column values lack corresponding parent table entries—through dialect-specific integrity check commands. "
        "Column and type validation operates independently, detecting duplicate column names within tables (case-insensitive comparison), "
        "unknown or unsupported data types, and tables with multiple primary key definitions. Data quality validation executes row count "
        "queries to identify empty tables—schema structures lacking actual data—which can confound training algorithms or produce "
        "misleading evaluation results when queries execute successfully against vacant schemas.",
    )

    add_paragraph(
        doc,
        "The architectural separation between structural validation (schema-level correctness) and data validation (row-level "
        "correctness) proves critical for remediation: structural errors require schema modifications and cannot be resolved through "
        "data corrections alone, whereas data violations often permit targeted cleaning operations without schema changes. Validation "
        "accuracy assessment through manual review of 89 flagged databases revealed 87.6% true positives (genuine integrity issues "
        "requiring remediation) and 12.4% false positives (constraints absent from formal declarations but enforced through application "
        "logic). The analyzer outputs comprehensive results including per-type error counts across all validation categories, severity "
        "classifications (Critical for errors blocking execution, High for integrity violations affecting correctness, Medium for design "
        "issues with limited impact), and detailed evidence (affected tables, columns, constraint definitions) enabling targeted "
        "debugging and repair workflows.",
    )

    # 2.4. Query Syntax Analyzer
    doc.add_page_break()
    add_heading(doc, "2.4. Query Syntax Analyzer", level=2)
    
    add_paragraph(
        doc,
        "The Query Syntax Analyzer performs static structural analysis of SQL statements, extracting syntactic features "
        "and computing complexity metrics without database execution. The analyzer processes queries through four stages: "
        "abstract syntax tree construction, structural feature extraction, complexity score calculation, and difficulty "
        "classification (Figure 2b).",
    )

    add_image_from_svg(
        doc,
        DIAGRAMS["fig2b"],
        "Figure 2b. Query Syntax Analyzer architecture with four-stage processing pipeline",
        width=6.5,
    )

    add_paragraph(
        doc,
        "The Abstract Syntax Tree Construction stage employs sqlglot, a SQL parser supporting over 20 dialects including SQLite, "
        "PostgreSQL, MySQL, and BigQuery. The parser transforms SQL text into a hierarchical tree where nodes represent syntactic "
        "constructs—SELECT expressions, JOIN operations, WHERE predicates, subqueries, and CTEs. This dialect-aware parsing is "
        "essential for datasets aggregating queries across multiple database systems, as constructs exhibit variations: string "
        "concatenation operators differ between dialects, and advanced features may be dialect-specific. Parsing failures are "
        "captured with detailed error messages facilitating identification of malformed queries.",
    )

    add_paragraph(
        doc,
        "The Structural Feature Extraction stage traverses the AST to count syntactic constructs relevant to complexity. The analyzer "
        "extracts five categories: Basic Structure (SELECT projections, WHERE predicates, GROUP BY, ORDER BY), Join and Relation "
        "(JOIN types, distinct tables, maximum depth), Advanced Features (subquery nesting, CTE complexity, window functions), "
        "Aggregation (COUNT, SUM, AVG, MIN, MAX, HAVING, DISTINCT), and Set Operations (UNION, INTERSECT, EXCEPT). This feature "
        "vector enables dataset composition analysis and correlation studies between syntactic patterns and semantic errors.",
    )

    add_paragraph(
        doc,
        "The Complexity Score Calculation stage synthesizes features into a 0-100 metric correlating with query difficulty. The algorithm "
        "implements weighted aggregation: single-table SELECT queries contribute 0-10 points, JOINs add 3-5 points, aggregates 2-4 points, "
        "and advanced features 5-10 points. Nesting depth penalties apply exponential scaling reflecting disproportionate cognitive costs, "
        "while well-structured CTEs and window functions receive bonuses. Raw scores undergo min-max normalization spanning 0-100.",
    )

    add_paragraph(
        doc,
        "The Difficulty Classification stage maps scores to four categorical labels with empirically calibrated thresholds. Easy queries "
        "(0-25) encompass single-table SELECT with simple predicates. Medium queries (26-50) introduce multi-table JOINs, GROUP BY, and "
        "straightforward subqueries requiring intermediate proficiency. Hard queries (51-75) exhibit multiple JOINs, nested subqueries, "
        "CTEs, and correlated subqueries challenging experienced practitioners. Expert queries (76-100) feature deep nesting, window "
        "functions, recursive CTEs, and intricate set operations representing frontier challenges. This enables benchmark coverage "
        "assessment, stratified evaluation, and imbalance identification.",
    )

    add_paragraph(
        doc,
        "The analyzer outputs structured metrics to the DuckDB database supporting dataset profiling, quality control, stratified evaluation, "
        "and correlation analyses exploring relationships between syntactic complexity and other quality dimensions.",
    )

    # 2.4.1. Query Antipattern Detector
    doc.add_page_break()
    add_heading(doc, "2.4.1. Query Antipattern Detector", level=2)
    
    add_paragraph(
        doc,
        "The Query Antipattern Detector implements rule-based static analysis identifying SQL code smells compromising maintainability, "
        "performance, correctness, or data safety. While the Syntax Analyzer characterizes complexity, the antipattern detector evaluates "
        "quality by detecting problematic patterns practitioners recognize as suboptimal or dangerous. The detector processes queries through "
        "four stages: AST traversal with pattern matching, severity classification, weighted scoring, and quality assignment (Figure 2c).",
    )

    add_image_from_svg(
        doc,
        DIAGRAMS["fig2c"],
        "Figure 2c. Query Antipattern Detector with four-stage quality assessment pipeline",
        width=6.5,
    )

    add_paragraph(
        doc,
        "The Rule-Based Detection Engine implements 14 independent rules traversing the AST to identify antipattern signatures. Each rule "
        "encapsulates domain knowledge about SQL best practices, performance optimization, and correctness pitfalls. Rules operate in parallel, "
        "enabling compositional analysis detecting multiple antipatterns simultaneously. The engine leverages sqlglot's AST node types: SELECT * "
        "detection searches for Star nodes; implicit join identifies FROM with multiple Tables but no Joins; function-in-WHERE locates Func nodes "
        "with Column references; leading wildcard LIKE examines patterns for wildcard prefixes; NOT IN combines Not wrapping In with Subqueries; "
        "correlated subquery performs scope analysis; unbounded query checks for missing Limits; unsafe UPDATE/DELETE identifies modifications "
        "lacking Where clauses; and additional rules target redundant DISTINCT with GROUP BY, unnecessary columns in EXISTS, UNION versus UNION ALL, "
        "complex OR chains, and DISTINCT overuse.",
    )

    add_paragraph(
        doc,
        "The Pattern Classification by Severity stage organizes antipatterns into three categories. Critical antipatterns (error severity) "
        "represent catastrophic patterns: unsafe UPDATE and DELETE lacking WHERE clauses unconditionally modifying or removing all rows. "
        "Significant antipatterns (warning severity) substantially compromise maintainability, performance, or correctness: SELECT * breaking "
        "interface contracts; implicit joins risking Cartesian products; functions in WHERE preventing index utilization; LIKE with leading "
        "wildcards; NOT IN with nullable subqueries; five or more JOINs; and DISTINCT on five or more columns. Advisory antipatterns (info "
        "severity) suggest optimization opportunities: correlated subqueries with N+1 patterns; unbounded SELECT queries; redundant DISTINCT "
        "with GROUP BY; unnecessary columns in EXISTS; UNION where UNION ALL suffices; and complex OR predicates.",
    )

    add_paragraph(
        doc,
        "The Quality Score Computation stage synthesizes antipatterns into a 0-100 metric through weighted penalties. Scoring begins at 100 points. "
        "Critical antipatterns subtract 20 points; significant antipatterns 10 points; advisory antipatterns 3 points. The formula: Quality Score = "
        "max(0, 100 - 20×critical - 10×warning - 3×info). This produces intuitive distributions: zero antipatterns achieve 100; advisory issues "
        "score 90s; several significant antipatterns score 50-70; critical antipatterns score below 50 necessitating revision or exclusion.",
    )

    add_paragraph(
        doc,
        "The Quality Level Classification stage maps scores to categorical labels with empirically calibrated thresholds. Excellent (90-100) "
        "designates professional-grade SQL with zero significant antipatterns. Good (70-89) encompasses one or two significant antipatterns "
        "representing acceptable quality. Fair (50-69) identifies multiple significant antipatterns requiring remediation. Poor (0-49) "
        "designates critical antipatterns or numerous issues requiring extensive revision or exclusion. This enables dataset-level profiling, "
        "quality-based filtering, correlation analyses with semantic correctness, and targeted improvement campaigns.",
    )

    add_paragraph(
        doc,
        "The detector outputs quality metrics to the DuckDB database enabling quality profiling, filtering, correlation studies with semantic "
        "correctness, targeted remediation campaigns, and quality-aware evaluation frameworks.",
    )

    # 2.5. Data flow
    doc.add_page_break()
    add_heading(doc, "2.5. Data Flow and Processing Model", level=2)

    add_image_from_svg(
        doc,
        DIAGRAMS["fig3"],
        "Figure 3. Data flow through the pipeline from input to output",
        width=6.5,
    )

    add_paragraph(
        doc,
        "Figure 3 illustrates the complete lifecycle of a single dataset record as it traverses the processing pipeline, emphasizing "
        "the streaming nature of the architecture. The processing model is fundamentally based on the iterator pattern, where each "
        "record flows through the pipeline stages sequentially without being accumulated in memory alongside other records. This "
        "design choice ensures O(1) memory complexity with respect to dataset size, enabling the system to process datasets containing "
        "millions of records on commodity hardware with modest RAM allocations.",
    )

    add_paragraph(
        doc,
        "The processing sequence begins when a configured loader reads a raw record from its data source, yielding a Python dictionary "
        "representation of that record's fields. This dictionary enters the normalization pipeline, where it passes sequentially through "
        "the three normalizers described previously. Each normalizer receives the record, performs its specific transformation or validation, "
        "and yields the modified record to the next normalizer. Upon completion of normalization, the record has been transformed into a "
        "standardized DataItem instance with guaranteed presence of all required fields and validated database identity.",
    )

    add_paragraph(
        doc,
        "The normalized DataItem then enters the analysis pipeline, implemented as a chain of generators where each analyzer acts as both "
        "a consumer of the preceding analyzer's output and a producer for the subsequent analyzer. Each analyzer receives the DataItem, "
        "executes its specific analysis logic (schema validation, syntax parsing, execution testing, antipattern detection, or semantic "
        "judgment), writes its resulting metrics to the DuckDB database through the configured MetricsSink, annotates the DataItem's "
        "metadata dictionary with summary information, and yields the enriched DataItem to the next analyzer in the chain. This generator-based "
        "architecture ensures that all analyses for a single record complete before the system begins processing the next record, maintaining "
        "strict streaming semantics and preventing memory accumulation.",
    )

    add_paragraph(
        doc,
        "Upon completion of all configured analyzers, the fully-annotated DataItem reaches the output layer, where two parallel operations "
        "occur. First, the annotated DataItem is serialized to the JSONL output file, preserving all original fields supplemented with "
        "the metadata accumulated from each analyzer, creating a self-contained record that can be filtered, sorted, or processed by "
        "downstream tools without requiring access to the metrics database. Second, the consolidated metrics from the DuckDB database "
        "reflect the cumulative results of all analyzers for this record, with batched writes ensuring efficient database performance "
        "despite the streaming processing model. This dual-output strategy provides both granular record-level access through the annotated "
        "dataset and efficient aggregate analysis through the structured metrics database.",
    )

    add_paragraph(
        doc,
        "Performance characteristics of the streaming pipeline exhibit distinct behavior depending on which analyzers are enabled. With "
        "only the four formal analyzers (schema validation, syntax analysis, execution testing, antipattern detection) active, typical "
        "processing rates range from 50 to 100 milliseconds per record on modern hardware, enabling complete analysis of thousand-record "
        "datasets within minutes. However, when the Semantic LLM Judge analyzer is enabled, processing time increases dramatically to "
        "approximately 2-10 seconds per record, dominated by network latency for API-based language model services and model inference "
        "time. For the Spider dataset with 2,147 test examples, this translates to roughly 12.7 hours of wall-clock time when using "
        "two concurrent language model judges (Gemini 2.5 Pro and GPT-4o), highlighting the computational cost of semantic validation "
        "and explaining why this analyzer is often selectively disabled for initial dataset exploration or when semantic validation is "
        "not required for the analysis objectives.",
    )

    # 2.6. DB Identity Normalizer
    doc.add_page_break()
    add_heading(doc, "2.6. Database Identity Normalization Algorithm", level=2)

    add_image_from_svg(
        doc,
        DIAGRAMS["fig5a"],
        "Figure 5a. DbIdentity normalization process: guaranteed dbId filling via health check or creation from DDL",
        width=6.5,
    )

    add_paragraph(
        doc,
        "The Database Identity Normalizer addresses a fundamental challenge in Text-to-SQL dataset processing: ensuring that every "
        "record can be associated with a valid, accessible database instance against which SQL queries can be validated and executed. "
        "This challenge arises from the diversity of dataset formats and conventions employed across the Text-to-SQL research community, "
        "where some datasets provide explicit database identifiers referencing pre-existing database files, others include DDL schemas "
        "allowing dynamic database construction, and some provide both, occasionally with inconsistencies between the identifier and schema.",
    )

    add_paragraph(
        doc,
        "The algorithm implements a decision tree with three primary branches corresponding to different combinations of database "
        "identifier presence and schema availability (Figure 5a). In the first scenario, when a record provides an explicit database "
        "identifier, the normalizer invokes a health check through the database manager, querying whether the referenced database exists "
        "and can be successfully opened. If the health check succeeds, the record passes through unchanged with the assumption that the "
        "provided database identifier is valid and all subsequent analyzers can safely reference it. If the health check fails but a DDL "
        "schema is present in the record, the normalizer attempts to recreate the database from the schema using the original identifier, "
        "effectively repairing a corrupted or missing database file. If the health check fails and no schema is available, the normalizer "
        "logs a warning but allows the record to proceed, acknowledging that while suboptimal, this situation may represent datasets where "
        "database files are expected to be provided separately from the dataset records.",
    )

    add_paragraph(
        doc,
        "The second scenario handles records lacking explicit database identifiers but including DDL schema specifications, a pattern "
        "commonly encountered in synthetic datasets or benchmark subsets designed for testing schema understanding capabilities. For these "
        "records, the normalizer computes a cryptographic hash (SHA-256) of the DDL schema text, generating a deterministic identifier "
        "that remains consistent across multiple pipeline executions processing the same schema. The normalizer then checks whether a "
        "database with this computed identifier already exists; if so, it reuses that database instance, leveraging caching to avoid "
        "redundant database creation operations when multiple records share identical schemas. If no such database exists, the normalizer "
        "creates a new database from the DDL schema, assigns it the computed hash identifier, and populates the record's dbId field with "
        "this identifier. This deterministic identifier generation strategy ensures reproducibility while supporting the analysis of "
        "DDL-only datasets that lack pre-existing database files.",
    )

    add_paragraph(
        doc,
        "The third scenario represents failure cases where records provide neither a database identifier nor a DDL schema. In this "
        "situation, the normalizer raises a ValueError with a descriptive error message, immediately halting processing of that record. "
        "This fail-fast behavior reflects a deliberate design decision: rather than attempting graceful degradation by skipping validation "
        "steps that require database access, the system treats the absence of database identification as a fatal data quality issue that "
        "must be addressed at the source. This strict validation ensures that downstream analyzers can rely on the guaranteed presence of "
        "a valid database identifier, simplifying their implementation and preventing silent failures where analyses would be skipped due "
        "to missing database access.",
    )

    add_paragraph(
        doc,
        "The normalization strategy supports several important use cases encountered in practical Text-to-SQL research. Classical datasets "
        "like Spider and BIRD, which provide explicit database files separate from the query annotations, benefit from the health check "
        "mechanism that validates database accessibility before attempting analysis, catching file corruption or missing files early in the "
        "pipeline. Synthetic DDL-only datasets, generated programmatically for testing purposes, utilize the hash-based identifier generation "
        "to enable analysis without requiring separate database file preparation. Mixed datasets containing heterogeneous records with different "
        "combinations of database identifiers and schemas can be processed uniformly through the branching logic that handles each case "
        "appropriately. The deterministic nature of the hash-based identifier generation ensures that pipeline re-execution produces identical "
        "database identifiers for identical schemas, facilitating reproducibility of experimental results and enabling caching strategies where "
        "database creation can be amortized across multiple pipeline runs.",
    )

    add_paragraph(
        doc,
        "Performance characteristics of database identity normalization vary significantly across the three scenarios. Health checks for "
        "existing, valid databases typically complete within 5-10 milliseconds, dominated by filesystem access to verify file existence and "
        "open the database file. Database recreation from DDL schemas when health checks fail requires 50-200 milliseconds, depending on "
        "schema complexity and the number of tables, indices, and foreign key constraints to be created. Initial database creation from DDL "
        "for new schemas exhibits similar timing characteristics. The ValueError exception raising for invalid records completes in under "
        "0.01 milliseconds. For the Spider test dataset with 2,147 records referencing 40 unique databases, the cumulative time spent in "
        "database identity normalization totals approximately 21 seconds, representing roughly 1% of the overall pipeline execution time "
        "when LLM-based semantic validation is disabled. This minimal overhead validates the design decision to perform health checks on "
        "every record rather than implementing a caching strategy, as the simplicity and fail-fast behavior of unconditional validation "
        "outweigh the modest performance cost.",
    )

    # 2.7. Protocol-based architecture
    doc.add_page_break()
    add_heading(doc, "2.7. Protocol-Based Architecture and Dependency Injection", level=2)

    add_image_from_svg(
        doc,
        DIAGRAMS["fig4"],
        "Figure 4. Protocol-based architecture with automatic dependency injection",
        width=6.5,
    )

    add_paragraph(
        doc,
        "The system architecture employs Python's Protocol typing construct to establish clear contracts between components while "
        "maintaining loose coupling between interface definitions and their concrete implementations (Figure 4). This architectural "
        "pattern, analogous to interface-based programming in statically typed languages, provides the benefits of explicit type "
        "checking and interface documentation while preserving the flexibility and dynamism characteristic of Python development.",
    )

    add_paragraph(
        doc,
        "The AnnotatingAnalyzer Protocol defines the essential contract that all analyzer implementations must satisfy, specifying "
        "required attributes including a string-valued name identifier and a boolean enabled flag for activation control, along with "
        "a mandatory analyze method signature accepting an iterable of DataItem instances, a MetricsSink for metrics persistence, and "
        "a string dataset identifier, while returning an iterator of potentially modified DataItem instances. This protocol enables "
        "the pipeline core to treat all analyzers uniformly regardless of their specific validation logic, facilitating the dynamic "
        "construction of analysis chains based on configuration file specifications without requiring central registry maintenance or "
        "explicit analyzer imports in the core pipeline logic.",
    )

    add_paragraph(
        doc,
        "The MetricsSink Protocol abstracts the mechanisms through which analyzers persist their computed metrics, defining three "
        "essential operations: write for accepting individual metric events, flush for ensuring buffered metrics are committed to "
        "persistent storage, and close for releasing any resources held by the sink. This abstraction enables multiple sink implementations "
        "with different performance and storage characteristics to coexist within the same system. The DuckDBSink implementation provides "
        "high-performance analytical access through SQL interfaces while automatically handling batched writes to amortize transaction "
        "overhead. The JsonlSink implementation offers a simple file-based storage format suitable for lightweight deployments or "
        "integration with external processing tools that expect line-delimited JSON. Additional sink implementations can be added without "
        "modifying analyzer code, as analyzers depend only on the protocol interface rather than specific sink implementations.",
    )

    add_paragraph(
        doc,
        "The Normalizer Protocol specifies the interface for components responsible for data standardization, requiring a normalize "
        "method that accepts an iterable of raw dictionary objects and returns an iterator of standardized DataItem instances. This "
        "protocol enables the composition of normalization pipelines where multiple normalizers operate in sequence, each consuming the "
        "output of the previous normalizer and producing input for the next. The loose coupling achieved through this protocol interface "
        "allows normalizers to be independently developed, tested, and versioned without impact on other pipeline components, facilitating "
        "the addition of new normalization strategies to handle emerging dataset formats or conventions without requiring modifications "
        "to the core pipeline infrastructure.",
    )

    add_paragraph(
        doc,
        "The DbAdapter Protocol encapsulates dialect-specific database operations, defining methods for query execution with "
        "configurable safety modes (read-only versus full read-write), schema introspection for retrieving table and column metadata, "
        "health checking for validating database accessibility, DDL export for extracting complete schema definitions, and database "
        "creation from DDL text specifications. This abstraction layer enables the system to support multiple SQL dialects (SQLite, "
        "PostgreSQL, MySQL) through different adapter implementations while presenting a uniform interface to higher-level components. "
        "The protocol-based approach facilitates unit testing by allowing test code to provide mock adapter implementations that simulate "
        "database operations without requiring actual database instances, significantly accelerating test suite execution and eliminating "
        "dependencies on external database installations.",
    )

    add_paragraph(
        doc,
        "Component orchestration within the protocol-based architecture relies on the dependency_injector library to automate the "
        "complex wiring of object graphs that arise from interdependent components. Each component declares its dependencies through "
        "a class-level INJECT attribute containing a list of dependency names. For example, the SchemaValidationAnalyzer specifies "
        "INJECT = [\"db_manager\"], indicating its requirement for a database manager instance. The dependency injection container, "
        "configured from YAML files, automatically instantiates dependencies, injects them into component constructors, and resolves "
        "transitive dependency chains. This automation eliminates substantial boilerplate code that would otherwise be required for "
        "manual object construction and wiring, while simultaneously improving testability by enabling test code to provide mock "
        "implementations for injected dependencies without modifying production code.",
    )

    add_paragraph(
        doc,
        "The protocol-based architecture with dependency injection offers several significant advantages for research software development. "
        "The explicit protocol interfaces serve as machine-checkable documentation of component contracts, with Python's runtime_checkable "
        "decorator enabling isinstance checks against protocol types for defensive programming patterns. Type checkers such as mypy can "
        "verify at static analysis time that implementations correctly satisfy their declared protocol obligations, catching interface "
        "violations before runtime. The loose coupling between interfaces and implementations facilitates independent evolution of components, "
        "allowing researchers to experiment with alternative implementations of specific analysis strategies without modifying surrounding "
        "infrastructure. The architecture naturally supports A/B testing scenarios where multiple implementations of the same protocol are "
        "evaluated under identical conditions. Finally, the separation of concerns between business logic (implemented in concrete classes) "
        "and composition logic (specified in configuration files) enables non-programmers to assemble analysis pipelines from pre-built "
        "components through declarative YAML configuration, lowering barriers to adoption and experimentation.",
    )

    # 2.8. LLM Judge Analyzer
    doc.add_page_break()
    add_heading(doc, "2.8. Multi-Model Semantic Validation Architecture", level=2)

    add_paragraph(
        doc,
        "The Semantic LLM Judge Analyzer represents the most sophisticated component of the validation pipeline, employing large language "
        "models to assess whether SQL queries correctly answer their associated natural language questions. This semantic validation capability "
        "addresses a critical gap in traditional formal validation approaches: while syntax checking and execution testing can verify technical "
        "correctness, only semantic analysis can determine whether a syntactically valid, successfully executing query actually implements the "
        "intended meaning of the natural language question. The analyzer implements this capability through a four-stage pipeline that combines "
        "intelligent context generation, prompt engineering, multi-model consensus voting, and result aggregation (Figure 6). The prompt structure "
        "with its four-category verdict classification system is detailed in Figure 6a.",
    )

    add_image_from_svg(
        doc,
        DIAGRAMS["fig6"],
        "Figure 6. Architecture of the Semantic LLM Judge with multi-model consensus voting",
        width=6.5,
    )

    add_paragraph(
        doc,
        "The first stage, Smart DDL Generation, addresses the token efficiency challenge inherent in providing database schema context to "
        "language models. The system supports two DDL generation modes: query_derived mode, which extracts only tables explicitly referenced "
        "in the SQL query to minimize token consumption, and full mode, which transmits complete database schemas for maximal context. The "
        "query_derived algorithm employs a multi-step optimization process. First, the system parses the SQL query using sqlglot to construct "
        "an abstract syntax tree, extracting the set of tables explicitly referenced in FROM clauses, JOIN specifications, and subqueries. "
        "Second, the database adapter queries system catalogs to retrieve metadata only for these referenced tables, including column names, "
        "data types, primary keys, and foreign key relationships. Third, the system executes SELECT queries to sample a configurable number "
        "of example values from each column (defaulting to 2 examples per column), providing concrete data instances that help language models "
        "understand the semantic content of columns beyond what data type information conveys. Finally, the collected metadata is formatted into "
        "CREATE TABLE statements with inline comments containing the example values. The query_derived mode typically reduces token consumption "
        "by factors of 5-10 compared to full mode, while the full mode provides complete database context for queries requiring comprehensive "
        "schema understanding or when table dependencies are complex.",
    )

    add_paragraph(
        doc,
        "The second stage, Prompt Template Resolution, combines the generated smart DDL with other contextual elements to produce a complete "
        "prompt for language model evaluation. The template, loaded from a YAML configuration file to facilitate experimentation with different "
        "prompt formulations, contains placeholders that are filled with dynamic values specific to each query being analyzed. The {{question}} "
        "placeholder receives the natural language question from the dataset record, providing the ground truth for what the SQL query should "
        "compute. The {{sql}} placeholder contains the SQL query under evaluation, formatted with syntax highlighting and whitespace normalization "
        "for improved readability. The {{ddl_schema}} placeholder is populated with the smart DDL generated in the first stage, supplying the "
        "database schema context necessary for semantic reasoning. An optional {{dialect}} placeholder specifies the SQL dialect (sqlite, "
        "postgresql, etc.), helping language models correctly interpret dialect-specific syntax variations. The assembled prompt specifies a "
        "strictly structured response format in which the model must return a single JSON object with exactly two keys, \"verdict\" and "
        "\"explanation\" (Figure 6a). The \"verdict\" field is constrained to one of four categorical labels: CORRECT (the query fully and accurately answers "
        "the question with no semantic issues), PARTIALLY_CORRECT (the query is mostly correct but exhibits minor deficiencies such as missing "
        "DISTINCT clauses when duplicates are possible, imprecise date/time boundaries, NULL-handling edge cases, or non-deterministic GROUP BY "
        "selections that do not fundamentally invalidate the result), INCORRECT (the query contains fundamental semantic errors including wrong "
        "table or column selections, missing or incorrect JOIN conditions, flawed aggregation logic, or erroneous WHERE clause predicates that "
        "produce fundamentally incorrect results), and UNANSWERABLE (the question cannot be answered from the provided schema due to missing "
        "tables or columns, requirements for external data not present in the database, or inherent ambiguity or contradictions in the question "
        "itself). The \"explanation\" field encodes a short natural-language justification; for CORRECT verdicts, the explanation is required to "
        "be an empty string, thereby preventing spurious narrative output, while for all other verdicts, the explanation must be a concise "
        "single-line description (bounded to at most 160–200 characters) that identifies the dominant semantic issue. This response specification "
        "transforms otherwise free-form LLM outputs into machine-parseable judgments that can be aggregated systematically across large datasets.",
    )

    # Figure 6a: Prompt Structure
    add_image_from_svg(
        doc,
        DIAGRAMS["fig6a"],
        "Figure 6a. LLM prompt structure with input parameters, evaluation criteria, and four-category verdict classification",
        width=6.5,
    )

    add_paragraph(
        doc,
        "The query_derived mode introduces methodological considerations requiring careful interpretation of validation results. When language "
        "models receive only query-referenced tables rather than complete schemas, UNANSWERABLE verdicts become ambiguous: models cannot "
        "distinguish between queries referencing semantically incorrect tables (which should be classified as INCORRECT when the full schema "
        "reveals that correct alternative tables exist) versus questions genuinely unanswerable due to inherent information gaps. This necessitates "
        "a two-stage validation protocol: queries receiving UNANSWERABLE verdicts in query_derived mode should be re-evaluated using full schema "
        "mode to disambiguate these scenarios. If re-evaluation yields INCORRECT, this indicates wrong table selection—a fundamental semantic error. "
        "If re-evaluation maintains UNANSWERABLE, this confirms genuine unanswerability. Additionally, the sampling of instance values, while "
        "providing semantic context about column content, may introduce bias when sampled values are unrepresentative of column distributions or "
        "mask data ambiguities that affect query correctness. Future work should include ablation studies isolating the effect of value sampling "
        "by comparing validation outcomes with and without sampled values, with varying sample sizes, and across query_derived versus full mode "
        "configurations, enabling quantification of how schema context completeness and instance value sampling independently influence semantic "
        "validation accuracy and verdict distribution patterns.",
    )

    add_paragraph(
        doc,
        "The third stage implements the multi-model consensus voting mechanism, querying multiple language models in parallel and collecting "
        "their independent judgments. The architecture supports four provider categories, each with different characteristics relevant to "
        "research applications. OpenAI models (GPT-4o, GPT-4-turbo, GPT-3.5-turbo) offer strong reasoning capabilities and reliable structured "
        "output generation, though with relatively high per-token costs. Anthropic models (Claude-3-opus, Claude-3-sonnet, Claude-3-haiku) "
        "provide alternative reasoning approaches that often exhibit different error patterns from GPT models, valuable for identifying edge "
        "cases where consensus breaks down. Google models (Gemini-2.5-pro, Gemini-1.5-pro, Gemini-1.5-flash) offer competitive performance "
        "with distinct training data and architectural choices that can surface different perspectives on ambiguous cases. Ollama models "
        "(Llama-3.1-70b, Mistral, Mixtral) enable fully local execution without API costs or rate limits, though typically with reduced "
        "accuracy compared to frontier commercial models. Each model receives the same prompt and returns a structured JSON response whose "
        "\"verdict\" component is one of the four allowed labels and whose \"explanation\" component follows the rules described above. The "
        "use of a shared, machine-enforced output schema ensures that model judgments are directly comparable across providers and model "
        "families, eliminating the need for heuristic parsing of free-form text. This multi-model design provides a natural defence against "
        "systematic biases of individual models: agreement across heterogeneous architectures increases trust in the underlying judgment, "
        "whereas systematic disagreement highlights records whose semantic interpretation is intrinsically ambiguous or whose annotation "
        "quality is suspect.",
    )

    add_paragraph(
        doc,
        "The fourth stage aggregates the individual model responses through weighted voting to determine consensus verdicts and identify "
        "cases requiring human review. Each model is associated with a non-negative scalar weight specified in the configuration file, allowing "
        "researchers to encode prior beliefs about the relative reliability or cost-effectiveness of different providers (for example, assigning "
        "higher weights to empirically validated frontier models while giving lower weights to cheaper or fully local models). The aggregation "
        "procedure operates along two complementary axes. First, it computes an unweighted majority verdict by counting, for each of the four "
        "categories, how many models produced that verdict; a majority consensus is deemed to exist when a single category accounts for more "
        "than half of all non-failed model responses, and a verdict is unanimous when all such responses coincide. Second, it derives a scalar "
        "semantic quality score in the closed interval [0,1] by mapping CORRECT to 1.0, PARTIALLY_CORRECT to 0.5, and both INCORRECT and "
        "UNANSWERABLE to 0.0, and then computing a provider-weighted average according to "
        "score = (∑_i w_i · v_i) / (∑_i w_i), where v_i denotes the mapped verdict value and w_i the corresponding model weight. This score "
        "captures the aggregate tendency of weighted models to regard a query as semantically adequate, while the majority and unanimity "
        "indicators characterize the degree of inter-model agreement. Records lacking any strict majority or exhibiting strongly divergent "
        "verdicts across models are of particular interest for dataset curation, as they frequently correspond to questions whose intent is "
        "underspecified, to SQL queries with subtle semantic flaws, or to cases where multiple, mutually incompatible interpretations of the "
        "question are plausible.",
    )

    add_paragraph(
        doc,
        "The multi-model semantic validation architecture addresses several research challenges in Text-to-SQL evaluation. By employing "
        "multiple models with different architectures and training data, the system reduces the risk of systematic biases that might arise "
        "from relying on a single model family. The consensus voting mechanism provides natural uncertainty quantification, with disagreement "
        "among models serving as a signal that examples may be problematic or ambiguous rather than clearly correct or incorrect. The "
        "structured output format with confidence scores enables downstream analysis to weight semantic validation results appropriately, "
        "for example by flagging low-confidence verdicts for human review or by analyzing whether certain query patterns systematically "
        "produce low-confidence judgments. The architecture's support for local models through Ollama integration enables researchers to "
        "trade accuracy for cost and throughput, running preliminary validation with local models before incurring API costs for expensive "
        "frontier models on the subset of examples requiring high-confidence judgments.",
    )

    # 2.9. Metrics system
    doc.add_page_break()
    add_heading(doc, "2.9. Two-Tier Metrics Architecture", level=2)

    add_paragraph(
        doc,
        "The metrics architecture implements a two-layer design that separates raw metric persistence from human-oriented presentation, "
        "reflecting the distinct requirements of these two use cases. The lower tier, mandatory in all pipeline configurations, handles "
        "efficient structured storage of analysis results in a format optimized for machine processing and complex analytical queries. "
        "The upper tier, optionally enabled through configuration, transforms the structured metrics into human-readable reports formatted "
        "in Markdown with explanatory text, formatted tables, and interpretation guidance.",
    )

    add_paragraph(
        doc,
        "DuckDB serves as the foundation of the metrics storage layer, selected for its combination of analytical query performance and "
        "operational simplicity. Unlike traditional OLAP databases requiring separate server processes and complex administration, DuckDB "
        "operates as an embedded database similar to SQLite, storing data in a single file that can be copied, versioned, and archived using "
        "standard filesystem operations. However, unlike SQLite's row-oriented storage optimized for transactional workloads, DuckDB employs "
        "columnar storage with aggressive compression, enabling analytical queries to process millions of metric records with sub-second "
        "response times even on laptop hardware. The system writes metrics to DuckDB using batched insert strategies that group hundreds or "
        "thousands of metric events into single transactions, amortizing the transaction commit overhead while maintaining streaming processing "
        "semantics at the record level through the use of in-memory buffering and periodic flush operations.",
    )

    add_paragraph(
        doc,
        "The structured metrics database contains five primary tables corresponding to the five analyzer types, with each table capturing "
        "the metrics specific to that analyzer's validation domain. Schema validation metrics record database identifiers, foreign key "
        "error categories and counts, duplicate column detections, and data type compatibility issues. Query syntax metrics store complexity "
        "scores, difficulty classifications, structural feature vectors (JOIN counts, subquery depths, CTE usage, aggregation functions, "
        "window functions), and parsing success/failure flags. Query execution metrics capture execution success/failure status, execution "
        "time in microseconds, returned row counts, and categorized error messages for failed executions. Antipattern detection metrics "
        "record individual antipattern flags for each of the 14 detected patterns, pattern severity scores, and overall quality scores "
        "computed from the weighted combination of detected patterns. Semantic validation metrics store individual model verdicts and "
        "confidence scores, consensus verdicts and types, weighted confidence aggregates, and the full text of model explanations for "
        "qualitative analysis. All tables share a common record identifier field enabling join operations that correlate metrics across "
        "different validation dimensions, facilitating research questions such as whether queries with higher syntactic complexity exhibit "
        "higher rates of semantic errors or whether certain antipattern combinations correlate with execution failures.",
    )

    add_paragraph(
        doc,
        "The separation between metrics persistence and report generation offers several operational advantages for research workflows. "
        "First, the metrics database serves as the canonical source of truth for all analysis results, providing unrestricted SQL query "
        "access for exploratory analysis beyond the predefined report templates. Researchers can write custom analytical queries to test "
        "hypotheses about dataset characteristics, compute statistical summaries over arbitrary subsets of records, or export metrics to "
        "statistical computing environments like R or Python for advanced analysis techniques. Second, report generation can be enabled or "
        "disabled independently of the analysis pipeline execution, allowing rapid iteration on large datasets where spending minutes "
        "generating reports would significantly extend pipeline runtime. During initial dataset exploration, researchers can disable reports "
        "entirely to minimize latency, then enable them in subsequent runs once the analysis strategy has stabilized. Third, reports can be "
        "regenerated at any time from the metrics database without re-executing the computationally expensive analysis pipeline, particularly "
        "valuable when semantic validation with LLM models contributes hours of processing time. Researchers can experiment with different "
        "report formats, presentation styles, or statistical aggregation strategies by simply re-running the report generation command against "
        "the existing metrics database, incurring only seconds of overhead rather than hours of re-analysis.",
    )

    add_paragraph(
        doc,
        "The report generation subsystem produces seven specialized report types, each addressing a different analytical perspective on "
        "dataset quality. The Summary Report provides a high-level overview synthesizing key statistics from all analyzers, including total "
        "record counts, overall success rates for each validation dimension, distributions of complexity and quality scores, and prominent "
        "error categories, offering researchers a quick assessment of dataset characteristics suitable for inclusion in papers or documentation. "
        "The Schema Validation Report enumerates databases with integrity violations, detailing specific foreign key errors, duplicate columns, "
        "and data type issues on a per-database basis, facilitating targeted database repair efforts. The LLM Judge Issues Report filters "
        "metrics to show only records judged INCORRECT or UNANSWERABLE by language models, listing the natural language questions, SQL queries, "
        "consensus verdicts, model-specific verdicts, and reasoning explanations, enabling manual inspection of potential false positives or "
        "interesting edge cases. The Query Execution Issues Report similarly filters for execution failures, grouping errors by category and "
        "providing representative examples to diagnose systematic problems such as missing database files or unsupported SQL dialect features. "
        "The Query Structure Profile Report analyzes the distribution of syntactic features across the dataset, presenting histograms of "
        "complexity scores, difficulty level proportions, feature usage frequencies (JOINs, subqueries, aggregations), and identifying outlier "
        "queries with unusually high complexity for closer inspection. The Table Coverage Report aggregates query references by table name, "
        "identifying which tables are heavily utilized in the dataset versus which tables appear rarely or never, informing decisions about "
        "dataset balance and potential augmentation needs. Finally, the Query Quality Report summarizes antipattern detection results, showing "
        "distributions of quality scores, frequencies of individual antipattern types, and examples of particularly low-quality queries that "
        "may warrant exclusion or remediation.",
    )

    add_paragraph(
        doc,
        "The two-tier architecture reflects a design philosophy that prioritizes structured data preservation over human-oriented presentation, "
        "acknowledging that research needs evolve and that the questions researchers wish to ask of their datasets often cannot be anticipated "
        "at system design time. By persisting all metrics in a queryable database with full SQL access, the system ensures that researchers "
        "retain the flexibility to perform ad-hoc analyses tailored to their specific research questions, supplementing or replacing the "
        "predefined reports with custom analytical code when project needs demand it. This architectural decision trades some initial usability "
        "(researchers must understand the metrics database schema to write custom queries) for long-term flexibility (arbitrary analytical "
        "questions can be addressed without system modification), a trade-off appropriate for research software where adaptability and "
        "extensibility outweigh the benefits of a more prescribed, but less flexible, reporting interface.",
    )

    # Save document
    print(f"\n💾 Saving document: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print(f"✅ Done! Document saved: {OUTPUT_FILE}")
    print("\n📄 Section 2 (Academic v2) contains:")
    print("   • 10 subsections (2.1 – 2.9, with 2.4.1)")
    print("   • 10 diagrams (Figures 2, 2a, 2b, 2c, 3, 4, 5a, 6, 6a)")
    print("   • Section 2.3 - Schema Validation Analyzer with Figure 2a")
    print("   • Section 2.4 - Query Syntax Analyzer with Figure 2b")
    print("   • Section 2.4.1 - Query Antipattern Detector with Figure 2c (NEW)")
    print("   • Section 2.6 - Database Identity Normalization (improved diagram)")
    print("   • Section 2.8 - Multi-Model Semantic Validation")
    print("   • Academic prose style with minimized bullet lists")
    print("   • Full English text optimized for journal publication")
    print("   • All analyzer methodologies fully documented")
    print("   • Structure: Architecture → Analyzers → Data Flow → Infrastructure")


if __name__ == "__main__":
    generate_section2_docx()

