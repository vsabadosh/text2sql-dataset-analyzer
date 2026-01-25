#!/usr/bin/env python3
"""
Generation of Results (Section 4), Discussion (Section 5), and Conclusion (Section 6) 
for e-Informatica Software Engineering Journal - Academic Standard.

This script generates properly structured sections following e-Informatica requirements:
- Section 4: Results (pure empirical findings without interpretation)
- Section 5: Discussion (interpretation, implications, threats to validity)
- Section 6: Conclusion (summary, future work, final remarks)
"""
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ModuleNotFoundError as e:
    raise SystemExit(
        "Missing dependency 'python-docx'. Install it with:\n"
        "  python3 -m pip install python-docx\n"
        "Then re-run this script."
    ) from e

# Base paths
BASE_DIR = Path(__file__).parent
OUTPUT_FILE = BASE_DIR / "Section4_5_6_Results_Discussion_Conclusion_v3.docx"


def add_heading(doc, text, level=1):
    """Add a heading with formatting."""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.name = "Times New Roman"
    heading.runs[0].font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    heading.runs[0].font.bold = True
    return heading


def add_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add a paragraph with formatting."""
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
    return para


def add_table_with_data(doc, headers, rows, caption=None):
    """Add a table with data and caption."""
    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Fill headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Format headers
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fill data
    for row_idx, row_data in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            cells[col_idx].text = str(cell_data)
            # Format data
            for paragraph in cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                # Center numeric data
                if col_idx > 0:  # First column left, others center
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add caption
    if caption:
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].font.size = Pt(10)
        caption_para.runs[0].font.italic = True
        caption_para.runs[0].font.name = "Times New Roman"

    return table


def generate_section4_results(doc):
    """
    Section 4: Results
    Pure empirical findings without interpretation or discussion.
    """
    print("📊 Generating Section 4: Results...")
    
    # Section 4 title
    title = doc.add_heading("4. RESULTS", level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True

    # Experimental Setup and Introduction
    add_heading(doc, "4.0. Experimental Setup", level=2)
    
    add_paragraph(
        doc,
        "This section reports empirical results from applying the validation framework (Section 2) to all 11,840 examples "
        "across Spider 1.0's three partitions: Test (2,147 examples, 40 databases), Dev (1,034 examples, 20 databases), and "
        "Train (8,659 examples, 146 databases). The dataset was obtained from the official Yale LILY repository [4] in January "
        "2025. The experimental infrastructure employed the technology stack described in Section 2.1, with SQLite 3.45.0 as "
        "the target dialect matching Spider's original database format. LLM-based semantic validation used OpenAI GPT-5 and "
        "Google Gemini 2.5 Pro (both accessed November 2024-January 2025). In practice, these reasoning-optimized models were "
        "used under provider-default reasoning behavior; while the pipeline configuration exposes a temperature parameter, it "
        "did not reliably enforce deterministic decoding in our runs. "
        "Experiments executed on a single workstation (macOS 14.x, Apple Silicon, 16GB RAM) with total pipeline runtime of "
        "67.8 hours (12.7hr Test, 6.1hr Dev, 49.0hr Train).",
    )
    
    add_paragraph(
        doc,
        "The analysis pipeline operated in the streaming architecture described in Section 2.2, with analyzer components "
        "configured per the methodology detailed in Sections 2.3-2.8. Schema validation executed once per unique database "
        "identifier (206 total) with cached metadata reused for queries sharing databases. Query-level analyzers operated "
        "independently on each SQL statement. Antipattern detection followed the SQLite-specific severity configuration "
        "(Section 2.6) with 19 pattern types and penalty weighting (Critical: -30, High: -15, Medium: -5, Low: -2). Execution "
        "testing used the query execution analyzer configured with mode=all (see Section 2.5), which executes SELECT statements "
        "with automatic LIMIT 1 injection when a LIMIT clause is absent, and tests UPDATE/DELETE/INSERT statements inside a rollback "
        "transaction while blocking destructive operations (e.g., DROP/TRUNCATE/ALTER). Semantic judging used two-model consensus "
        "with verdict categories (CORRECT/PARTIALLY_CORRECT/INCORRECT/UNANSWERABLE) as defined in Section 2.8; configuration parameters "
        "included schema_mode=full, num_examples=2, and parallel voting enabled (parallel_voters=true, max_workers=2).",
    )
    
    add_paragraph(
        doc,
        "Results are organized across five quality dimensions detailed in Sections 2.3-2.8: (i) database schema integrity "
        "(Section 4.1); (ii) SQL syntactic structure (Section 4.2); (iii) dynamic executability (Section 4.3); (iv) antipattern "
        "prevalence (Section 4.4); (v) semantic correctness (Section 4.5). Sections 4.6-4.7 report computational performance and "
        "integrated quality assessment. Findings are presented descriptively without interpretation; Section 5 provides analysis "
        "of implications and validity threats.",
    )

    # 4.1. Schema Integrity
    doc.add_page_break()
    add_heading(doc, "4.1. Database Schema Integrity", level=2)
    
    add_paragraph(
        doc,
        "Table 1 presents comparative schema validation results across all three Spider partitions. Schema validation "
        "operates at the database level (one analysis per unique database identifier), examining structural integrity "
        "(foreign key definitions, column uniqueness, type validity) and data-level integrity (referential constraint "
        "violations in stored rows). The 'Valid databases' metric denotes databases passing all validation checks with "
        "no errors or warnings.",
    )

    # Table 1 (comparative schema results)
    headers_t1 = ["Metric", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t1 = [
        ["Total databases", "40", "20", "146"],
        ["Valid databases", "36 (90.0%)", "13 (65.0%)", "110 (75.3%)"],
        ["Databases with errors (non-warning)", "4 (10.0%)", "5 (25.0%)", "32 (21.9%)"],
        ["Databases with warnings only", "0 (0.0%)", "2 (10.0%)", "4 (2.7%)"],
        ["Non-clean databases (errors + warnings)", "4 (10.0%)", "7 (35.0%)", "36 (24.7%)"],
        ["Total tables", "180", "80", "793"],
        ["Empty tables", "0 (0%)", "0 (0%)", "71 (9.0%)"],
        ["Total foreign keys", "146", "64", "717"],
        ["Invalid foreign keys", "4", "7", "78"],
        ["Databases with FK data violations", "1", "3", "6"],
    ]
    add_table_with_data(
        doc,
        headers_t1,
        rows_t1,
        "Table 1. Comparative database schema validation results across Spider partitions",
    )

    add_paragraph(
        doc,
        "Table 1 reveals a quality gradient inverse to typical expectations: Test achieves 90.0% schema validity (36/40 "
        "databases clean), Train achieves 75.3% (110/146), while Dev—the partition used for model selection in standard "
        "workflows—exhibits the lowest validity at 65.0% (only 13/20 databases pass validation). Structural foreign key "
        "errors total 89 across all partitions (4 in Test, 7 in Dev, 78 in Train). The 'Empty tables' row shows Train "
        "contains 71 tables with zero rows (9.0% of Train's 793 tables), while Test and Dev have no empty tables. "
        "Referential integrity violations at the data level affect 10 databases total: 1 in Test, 3 in Dev, and 6 in Train.",
    )

    # Table 2 (FK error types)
    add_paragraph(
        doc,
        "Table 2 presents the distribution of foreign key error types across partitions, separating structural errors "
        "(schema/DDL-level definition problems, measured as count of invalid foreign key relationships) from data "
        "violations (referential integrity breaches in stored rows, measured as count of violating rows). The 'Severity' "
        "column reflects the potential impact: Critical errors always produce incorrect query results or data safety issues, "
        "High-severity errors cause wrong results in specific cases, while Medium/Low issues affect performance or style.",
    )

    headers_t2_fk = ["Error Type", "Spider Test", "Spider Dev", "Spider Train", "Total", "Severity"]
    rows_t2_fk = [
        ["STRUCTURAL ERRORS", "", "", "", "", ""],
        ["FK Missing Table", "0", "0", "0", "0", "Critical"],
        ["FK Missing Column", "3", "0", "5", "8", "Critical"],
        ["FK Arity Mismatch", "0", "0", "0", "0", "Critical"],
        ["FK Type Mismatch", "1", "5", "32", "38", "High"],
        ["FK Target Not Key", "0", "2", "41", "43", "High"],
        ["Duplicate Columns", "0", "0", "0", "0", "Medium"],
        ["Unknown Data Types", "0", "0", "0", "0", "Low"],
        ["Subtotal (Structural)", "4", "7", "78", "89", "—"],
        ["DATA VIOLATIONS", "", "", "", "", ""],
        ["FK Data Violations (rows)", "1", "2,403", "39,523", "41,927", "Critical"],
        ["Databases with Data Violations", "1", "3", "6", "10", "—"],
    ]
    add_table_with_data(
        doc,
        headers_t2_fk,
        rows_t2_fk,
        "Table 2. Foreign key error distribution by type and partition",
    )

    add_paragraph(
        doc,
        "Within structural errors, FK Type Mismatch (38 occurrences) and FK Target Not Key (43 occurrences) dominate, "
        "together accounting for 81 of 89 total structural FK errors (91.0%). FK Type Mismatch occurs when local and "
        "parent columns have incompatible SQL types or type affinities (e.g., INTEGER vs TEXT), preventing reliable "
        "join operations. FK Target Not Key indicates foreign keys referencing columns that are neither PRIMARY KEY nor "
        "UNIQUE, violating relational integrity principles that foreign keys must reference candidate keys. FK Missing "
        "Column (8 occurrences, all Critical severity) represents cases where foreign keys reference non-existent columns "
        "in the parent table. For data-level integrity, referential violations total 41,927 rows across 10 databases, "
        "with the vast majority concentrated in Train (39,523 rows) and notably including 2,403 violations in Dev.",
    )

    # Table 3 (FK Error Details by Database ID)
    add_paragraph(
        doc,
        "Table 3 enumerates specific database identifiers associated with each foreign key structural error type, "
        "enabling targeted remediation. The '×' notation indicates multiple occurrences within a single database "
        "(e.g., 'book_1 (2×)' means two FK Missing Column errors in book_1). This drill-down complements Table 2's "
        "aggregate counts by identifying which databases require repair and the multiplicity of issues per database.",
    )
    
    headers_t3_detail = ["Error Type", "Partition", "Count", "Database IDs"]
    rows_t3_detail = [
        ["FK Missing Column", "Test", "3", "book_1 (2×), car_racing (1×)"],
        ["FK Missing Column", "Dev", "0", "—"],
        ["FK Missing Column", "Train", "5", "baseball_1, imdb, loan_1, restaurants, store_product"],
        ["", "", "", ""],
        ["FK Type Mismatch", "Test", "1", "government_shift"],
        ["FK Type Mismatch", "Dev", "5", "car_1, concert_singer (2×), employee_hire_evaluation, museum_visit"],
        ["FK Type Mismatch", "Train", "32", 
         "academic, aircraft (2×), architecture (2×), city_record, company_employee, "
         "cre_Drama_Workshop_Groups (9×), culture_company (2×), loan_1, machine_repair, "
         "party_people, performance_attendance (2×), phone_1, phone_market, race_track, "
         "school_finance (2×), shop_membership (2×), student_assessment, wrestler"],
        ["", "", "", ""],
        ["FK Target Not Key", "Test", "0", "—"],
        ["FK Target Not Key", "Dev", "2", "car_1, voter_1"],
        ["FK Target Not Key", "Train", "41", 
         "baseball_1 (20×), dorm_1 (3×), imdb (6×), soccer_1 (4×), wine_1 (2×), yelp (7×)"],
    ]
    add_table_with_data(
        doc,
        headers_t3_detail,
        rows_t3_detail,
        "Table 3. Database identifiers by foreign key error type (× denotes error count per database)",
    )

    add_paragraph(
        doc,
        "FK Type Mismatch errors affect 19 distinct databases across partitions, with concentration in Train. Notably, "
        "cre_Drama_Workshop_Groups contains 9 instances (the highest multiplicity for any single database-error combination), "
        "suggesting systematic schema design issues in that database. FK Target Not Key errors exhibit even higher "
        "concentration: baseball_1 contains 20 occurrences, yelp contains 7, and imdb contains 6. These three databases "
        "alone account for 33 of 43 total FK Target Not Key errors (76.7%), indicating localized rather than distributed "
        "problems. FK Missing Column errors appear in 8 databases total (3 Test, 0 Dev, 5 Train), representing Critical "
        "severity issues where queries attempting JOINs on these foreign keys will fail execution due to non-existent columns.",
    )

    # Table 4 (FK violations by database)
    add_paragraph(
        doc,
        "Table 4 enumerates databases with foreign key data violations (referential integrity breaches in stored rows "
        "rather than schema definition errors). The '% of Total' column shows each database's contribution to the "
        "overall pool of 41,927 violating rows. This table reveals extreme concentration: the top two databases account "
        "for the vast majority of violations, while the remaining eight databases contribute negligibly.",
    )
    
    headers_t4 = ["Database ID", "Partition", "FK Data Violations", "% of Total"]
    rows_t4 = [
        ["sakila_1", "Train", "38,273", "91.3%"],
        ["flight_2", "Dev", "2,400", "5.7%"],
        ["flight_4", "Train", "1,240", "3.0%"],
        ["car_1", "Dev", "2", "<0.1%"],
        ["hr_1", "Train", "6", "<0.1%"],
        ["college_1", "Train", "2", "<0.1%"],
        ["wta_1", "Dev", "1", "<0.1%"],
        ["hospital_1", "Train", "1", "<0.1%"],
        ["allergy_1", "Train", "1", "<0.1%"],
        ["pilot_1", "Test", "1", "<0.1%"],
        ["TOTAL", "—", "41,927", "100%"],
    ]
    add_table_with_data(
        doc,
        headers_t4,
        rows_t4,
        "Table 4. Databases with foreign key data violations",
    )

    add_paragraph(
        doc,
        "Table 4 demonstrates catastrophic concentration: sakila_1 (Train) and flight_2 (Dev) together account for "
        "40,673 of 41,927 violations (97.0%). Specifically, sakila_1 contributes 38,273 violating rows (91.3% of total), "
        "representing a single database responsible for the vast majority of referential integrity failures across the "
        "entire benchmark. Flight_2 in Dev contributes 2,400 violations (5.7%), which is particularly concerning given "
        "Dev's role in model selection—these violations affect queries for approximately 35.6% of Dev examples sharing "
        "this database. The remaining eight databases collectively contain only 14 violations (<0.1%), with individual "
        "counts ranging from 1 to 6 rows. This distribution indicates that referential integrity problems are highly "
        "localized to two catastrophic cases rather than broadly distributed across Spider's databases.",
    )

    # Table 5 (Empty tables)
    add_paragraph(
        doc,
        "Table 5 identifies databases with empty tables (tables containing zero rows despite having defined schemas), "
        "found exclusively in the Train partition. 'Data Coverage' represents the percentage of tables containing at "
        "least one row. Databases with 0% coverage have all tables empty, rendering execution-based validation impossible "
        "for queries referencing those databases.",
    )
    
    headers_t5 = ["Database ID", "Total Tables", "Empty Tables", "Data Coverage"]
    rows_t5 = [
        ["academic", "15", "15", "0%"],
        ["imdb", "16", "16", "0%"],
        ["yelp", "7", "7", "0%"],
        ["restaurants", "3", "3", "0%"],
        ["geo", "7", "7", "0%"],
        ["music_2", "7", "7", "0%"],
        ["scholar", "10", "10", "0%"],
        ["sakila_1", "16", "4", "75%"],
        ["formula_1", "13", "2", "85%"],
        ["TOTAL (Train only)", "94", "71", "24.5%"],
    ]
    add_table_with_data(
        doc,
        headers_t5,
        rows_t5,
        "Table 5. Databases with empty tables in Spider Train partition",
    )

    add_paragraph(
        doc,
        "Seven Train databases exhibit complete data absence (0% coverage), accounting for 65 of the 71 empty tables: "
        "academic (15 empty tables), imdb (16), yelp (7), restaurants (3), geo (7), music_2 (7), and scholar (10). These "
        "databases retain only schema definitions (CREATE TABLE statements with column specifications) but contain no "
        "INSERT statements or populated rows, rendering any query execution against them meaningless for result validation. "
        "Two additional databases show partial emptiness: sakila_1 has 4 empty tables out of 16 (75% coverage), and "
        "formula_1 has 2 empty tables out of 13 (85% coverage). The Train partition total of 71 empty tables represents "
        "9.0% of Train's 793 tables. Test and Dev contain no empty tables, maintaining 100% data coverage across all "
        "tables in those partitions. Notably, three databases appear in both empty-table and foreign-key-error lists: "
        "imdb (16 empty tables + FK errors), yelp (7 empty tables + FK errors), and academic (15 empty tables + FK errors).",
    )

    # 4.2. Syntactic Structure
    doc.add_page_break()
    add_heading(doc, "4.2. SQL Syntactic Structure and Complexity", level=2)
    
    add_paragraph(
        doc,
        "Table 6 presents overall syntactic analysis results from the query-level SQL parser operating on all 11,840 "
        "SQL statements. 'Successfully parsed' indicates queries that produced valid abstract syntax trees (ASTs) using "
        "sqlglot's SQLite dialect parser, enabling feature extraction (tables, JOINs, subqueries, aggregations). "
        "'Complexity score' is a 0–100 structural metric computed via rule-based difficulty classification with "
        "micro-adjustments based on feature counts; mean, median, and standard deviation characterize the distribution.",
    )

    # Table 6 (syntactic analysis)
    headers_t6 = ["Metric", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t6 = [
        ["Total queries", "2,147", "1,034", "8,659"],
        ["Successfully parsed", "2,147 (100%)", "1,034 (100%)", "8,659 (100%)"],
        ["Parse failures", "0 (0%)", "0 (0%)", "0 (0%)"],
        ["Mean complexity score", "39.5", "39.0", "40.7"],
        ["Median complexity score", "41", "41", "41"],
        ["Standard deviation", "20.2", "19.9", "21.2"],
    ]
    add_table_with_data(
        doc,
        headers_t6,
        rows_t6,
        "Table 6. Overall syntactic analysis results",
    )

    add_paragraph(
        doc,
        "All 11,840 queries parsed successfully (100% across all partitions: 2,147/2,147 Test, 1,034/1,034 Dev, "
        "8,659/8,659 Train), indicating that Spider contains no unparseable or syntactically malformed SQL under "
        "SQLite dialect rules. Complexity scores exhibit tight clustering: mean values range narrowly from 39.0 (Dev) "
        "to 40.7 (Train), with identical median of 41 across all three partitions. This consistency suggests Test and "
        "Dev were sampled to maintain similar difficulty distributions to Train. Standard deviations (19.9–21.2) indicate "
        "moderate spread, with approximately 68% of queries scoring between 18–62 (mean ± 1 SD). The tight central "
        "tendency shows Spider's bulk queries concentrate in the medium-difficulty range (complexity 30–59), with smaller "
        "tails representing easy (10–29) and hard/expert (60–100) queries.",
    )

    # Table 7 (difficulty distribution)
    add_paragraph(
        doc,
        "Table 7 presents the distribution of queries by analyzer-assigned difficulty category (Easy/Medium/Hard/Expert) "
        "derived from structural complexity rules. Difficulty classification considers table count, JOIN count, presence "
        "of subqueries/CTEs/window functions, aggregation patterns, and set operations. The 'Complexity Range' column "
        "shows the approximate score ranges associated with each difficulty band in the summary (note: actual classification "
        "is rule-based on features, not purely score-threshold-based).",
    )

    headers_t7 = ["Difficulty Level", "Spider Test", "Spider Dev", "Spider Train", "Complexity Range"]
    rows_t7 = [
        ["Easy", "516 (24.0%)", "250 (24.2%)", "2,039 (23.5%)", "10-29"],
        ["Medium", "1,246 (58.0%)", "612 (59.2%)", "4,827 (55.7%)", "30-59"],
        ["Hard", "385 (17.9%)", "172 (16.6%)", "1,727 (19.9%)", "60-79"],
        ["Expert", "0 (0.0%)", "0 (0.0%)", "66 (0.8%)", "80-100"],
    ]
    add_table_with_data(
        doc, headers_t7, rows_t7, "Table 7. Query difficulty distribution across partitions"
    )

    add_paragraph(
        doc,
        "Difficulty distributions are remarkably consistent across Test and Dev: Easy accounts for approximately 24% "
        "(516/2,147 Test; 250/1,034 Dev), Medium for 58–59% (1,246/2,147 Test; 612/1,034 Dev), and Hard for ~17% "
        "(385/2,147 Test; 172/1,034 Dev). Train exhibits similar proportions for the first three categories (23.5% Easy, "
        "55.7% Medium, 19.9% Hard), confirming that Test and Dev sampling preserved Train's difficulty distribution. "
        "However, Train uniquely contains 66 Expert-level queries (0.8% of Train), featuring rare advanced patterns like "
        "recursive CTEs, deeply nested subqueries (depth ≥2), or complex combinations of multiple advanced features. "
        "Neither Test nor Dev contains Expert queries (0.0%), likely due to their smaller sample sizes (2,147 and 1,034 "
        "examples respectively) making it statistically unlikely to sample from the rare 0.8% Expert population. The "
        "dominance of Medium difficulty (55.7–59.2% across partitions) aligns with Spider's design goal of representing "
        "realistic cross-domain queries requiring JOINs, aggregations, and GROUP BY operations.",
    )

    # Table 8 (JOIN distribution)
    add_paragraph(
        doc,
        "Table 8 presents the distribution of queries by count of explicit `JOIN` operators detected in the parsed SQL "
        "abstract syntax tree. A JOIN count of 0 indicates queries containing no explicit `JOIN` clauses (INNER JOIN, "
        "LEFT JOIN, etc.); this category includes genuine single-table queries as well as multi-table queries expressed "
        "via comma-separated FROM lists (e.g., 'FROM users, orders') which do not produce explicit JOIN AST nodes. "
        "Categories 1 JOIN through 4+ JOINs count queries by the number of explicit JOIN keywords present. The final row "
        "'Multi-table (≥1 JOIN)' aggregates all queries with at least one explicit JOIN operator.",
    )

    headers_t8 = ["JOIN Count", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t8 = [
        ["0 (single-table)", "1,285 (59.85%)", "626 (60.54%)", "4,847 (55.98%)"],
        ["1 JOIN", "558 (25.99%)", "320 (30.95%)", "2,233 (25.79%)"],
        ["2 JOINs", "242 (11.27%)", "72 (6.96%)", "1,042 (12.03%)"],
        ["3 JOINs", "38 (1.77%)", "10 (0.97%)", "359 (4.15%)"],
        ["4+ JOINs", "24 (1.12%)", "6 (0.58%)", "178 (2.06%)"],
        ["Multi-table (≥1 JOIN)", "862 (40.15%)", "408 (39.46%)", "3,812 (44.02%)"],
    ]
    add_table_with_data(
        doc,
        headers_t8,
        rows_t8,
        "Table 8. JOIN distribution analysis",
    )

    add_paragraph(
        doc,
        "Single-table queries (0 explicit JOINs) dominate all partitions: 1,285/2,147 (59.85%) in Test, 626/1,034 (60.54%) "
        "in Dev, and 4,847/8,659 (55.98%) in Train. This majority reflects Spider's inclusion of substantial simple-query "
        "representation alongside cross-table queries. Conversely, queries with at least one explicit JOIN ('Multi-table') "
        "account for 40.15% (Test), 39.46% (Dev), and 44.02% (Train), with Train exhibiting slightly higher multi-table "
        "representation. Among multi-JOIN queries, 1 JOIN is most common (25.79–30.95%), followed by 2 JOINs (6.96–12.03%). "
        "Queries with 3 JOINs occur at frequencies ranging from 0.97% (Dev) to 4.15% (Train), while 4+ JOINs are rare "
        "(0.58–2.06%). Train consistently shows elevated frequencies in higher JOIN-count categories (3 JOINs: 4.15% vs "
        "1.77% Test; 4+ JOINs: 2.06% vs 1.12% Test), likely reflecting its larger sample size capturing more diverse query "
        "patterns. The 40–44% multi-table representation across partitions aligns with Spider's cross-domain focus, where "
        "realistic information needs often require joining multiple related tables.",
    )

    # 4.3. Query Execution
    doc.add_page_break()
    add_heading(doc, "4.3. Dynamic Query Executability", level=2)
    
    add_paragraph(
        doc,
        "Table 9 presents query execution results under the framework's safe execution protocol. Each SQL statement is "
        "executed against its associated database using SQLite 3.x, with SELECT queries automatically receiving LIMIT "
        "injection when absent (default LIMIT 1 for safety). The protocol validates that statements can be parsed and "
        "executed by the engine without requiring full result materialization. 'Mean execution time' measures end-to-end "
        "latency including connection overhead, query submission, and result fetch under the LIMIT-constrained protocol.",
    )

    # Table 9 (execution results)
    headers_t9 = ["Metric", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t9 = [
        ["Total queries", "2,147", "1,034", "8,659"],
        ["Successfully executed", "2,147 (100%)", "1,034 (100%)", "8,656 (99.97%)"],
        ["Execution failures", "0 (0%)", "0 (0%)", "3 (0.03%)"],
        ["Skipped queries", "0 (0%)", "0 (0%)", "0 (0%)"],
        ["Mean execution time (ms)", "3.10", "2.49", "1.19"],
    ]
    add_table_with_data(
        doc,
        headers_t9,
        rows_t9,
        "Table 9. Query execution results",
    )

    add_paragraph(
        doc,
        "Test and Dev achieve perfect execution success (100%: 2,147/2,147 and 1,034/1,034 respectively), with zero "
        "execution failures. Train achieves 99.97% success (8,656/8,659), with only 3 execution failures out of 8,659 "
        "queries (0.03%). Mean execution times remain remarkably low across all partitions: 3.10 ms (Test), 2.49 ms (Dev), "
        "and 1.19 ms (Train). The faster Train execution time likely reflects higher proportions of single-table queries "
        "(55.98% vs 59.85–60.54% in Test/Dev from Table 8) and the LIMIT-constrained protocol that caps result set sizes. "
        "The near-perfect executability (99.97–100%) demonstrates that Spider's SQL statements are syntactically and "
        "semantically well-formed at the execution level, capable of running against their associated databases under "
        "SQLite's permissive execution model. No queries are listed as 'Skipped', indicating that the execution analyzer "
        "attempted all 11,840 queries without pipeline-level skipping due to prior analyzer failures.",
    )

    # 4.4. Antipattern Detection
    doc.add_page_break()
    add_heading(doc, "4.4. SQL Code Quality and Antipatterns", level=2)
    
    add_paragraph(
        doc,
        "Table 10 presents overall code quality metrics from the SQL antipattern detector, which analyzes queries for "
        "patterns known to cause correctness issues, performance problems, or maintainability concerns. 'Mean quality score "
        "(/100)' aggregates per-query scores where detected antipatterns reduce the score from an ideal 100 based on "
        "severity-weighted penalties (Critical: -30, High: -15, Medium: -5, Low: -2). 'Queries without antipatterns' "
        "represents the fraction of queries passing all configured antipattern checks with no flagged issues.",
    )

    # Table 10 (antipattern overall quality)
    headers_t10 = ["Metric", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t10 = [
        ["Mean quality score (/100)", "97.7", "97.7", "98.4"],
        ["Average antipatterns per query", "0.2", "0.2", "0.1"],
        ["Queries without antipatterns", "1,772 (82.5%)", "865 (83.7%)", "7,494 (86.5%)"],
        ["Total antipattern occurrences", "383", "169", "1,182"],
    ]
    add_table_with_data(
        doc,
        headers_t10,
        rows_t10,
        "Table 10. Overall code quality metrics",
    )

    add_paragraph(
        doc,
        "Mean quality scores are consistently high across partitions: 97.7/100 (Test), 97.7/100 (Dev), and 98.4/100 (Train), "
        "indicating that most queries avoid severe antipatterns. The majority of queries contain no detected antipatterns: "
        "1,772/2,147 (82.5%) in Test, 865/1,034 (83.7%) in Dev, and 7,494/8,659 (86.5%) in Train. These clean-query "
        "percentages demonstrate that approximately 83–87% of Spider passes code quality checks under the configured "
        "severity thresholds. Conversely, 375 queries (17.5%) in Test, 169 (16.3%) in Dev, and 1,165 (13.5%) in Train "
        "contain at least one flagged antipattern, totaling 1,709 queries with antipatterns across the benchmark (14.4% "
        "of 11,840). Total antipattern occurrences sum to 1,734 (383 Test + 169 Dev + 1,182 Train), with the difference "
        "between 1,709 queries-with-antipatterns and 1,734 occurrences indicating that some queries contain multiple "
        "antipatterns simultaneously. Average antipatterns per query (0.2 Test/Dev, 0.1 Train) reflect the rarity of "
        "antipatterns relative to dataset size, with most occurrences concentrated in specific pattern types as detailed in Table 11.",
    )

    # Table 11 (antipattern distribution)
    add_paragraph(
        doc,
        "Table 11 presents antipattern distribution by type and configured severity level. Counts report the number of "
        "queries containing each antipattern type, with percentages computed relative to the partition's total query count. "
        "Severity classifications (Critical/High/Medium/Low) reflect configured detection rules: Critical patterns always "
        "produce wrong results, High patterns cause correctness failures in specific cases, Medium patterns create performance "
        "issues, and Low patterns represent style preferences acceptable in research datasets.",
    )

    headers_t11 = ["Antipattern", "Spider Test", "Spider Dev", "Spider Train", "Severity"]
    rows_t11 = [
        ["Cartesian product", "2 (0.1%)", "2 (0.2%)", "13 (0.2%)", "Critical"],
        ["Missing GROUP BY", "239 (11.1%)", "106 (10.3%)", "600 (6.9%)", "High"],
        ["NOT IN with nullable columns", "74 (3.4%)", "46 (4.4%)", "228 (2.6%)", "High"],
        ["Leading wildcard LIKE", "50 (2.3%)", "12 (1.2%)", "140 (1.6%)", "Medium"],
        ["Function in WHERE clause", "0 (0.0%)", "0 (0.0%)", "6 (0.1%)", "Medium"],
        ["Redundant DISTINCT", "6 (0.3%)", "0 (0.0%)", "133 (1.5%)", "Low"],
        ["SELECT *", "12 (0.6%)", "3 (0.3%)", "62 (0.7%)", "Low"],
    ]
    add_table_with_data(
        doc, headers_t11, rows_t11, "Table 11. Antipattern distribution by severity"
    )

    add_paragraph(
        doc,
        "Missing GROUP BY is the most frequent high-severity antipattern, totaling 945 occurrences across all partitions: "
        "239/2,147 (11.1%) in Test, 106/1,034 (10.3%) in Dev, and 600/8,659 (6.9%) in Train. This pattern occurs when "
        "queries use aggregate functions (COUNT, SUM, AVG, etc.) alongside non-aggregated columns without proper GROUP BY "
        "clauses, exploiting SQLite's permissive handling that would produce errors in stricter systems like PostgreSQL "
        "or MySQL with ONLY_FULL_GROUP_BY mode. NOT IN with nullable columns affects 348 queries total (74 Test (3.4%), "
        "46 Dev (4.4%), 228 Train (2.6%)), introducing subtle correctness risks under SQL's three-valued logic where NULL "
        "in subquery results causes the entire NOT IN predicate to evaluate to UNKNOWN rather than TRUE/FALSE. Critical "
        "Cartesian product detections total 17 queries across all partitions (2 Test, 2 Dev, 13 Train), representing "
        "queries with missing JOIN conditions that produce exponential result set explosions. All 17 Cartesian product "
        "cases underwent manual verification, confirming them as incorrect queries with broken join logic; notably, these "
        "same 17 queries received unanimous INCORRECT judgments from both LLM judges (Gemini 2.5 Pro and GPT-5) in semantic "
        "validation (Section 4.5), providing independent confirmation of their defect status. Medium-severity patterns "
        "include Leading wildcard LIKE (202 occurrences), Function in WHERE clause (6 occurrences), and low-severity "
        "patterns include Redundant DISTINCT (139 occurrences) and SELECT * (77 occurrences).",
    )

    # 4.5. Semantic Validation
    doc.add_page_break()
    add_heading(doc, "4.5. Semantic Correctness Assessment", level=2)
    
    add_paragraph(
        doc,
        "Table 12 presents semantic validation results obtained via two-model LLM consensus judging (Gemini 2.5 Pro and "
        "GPT-5). Each of the 11,840 queries received independent verdicts from both models, with "
        "possible verdicts being CORRECT (query semantically answers the question), PARTIALLY_CORRECT (query mostly correct "
        "with minor issues), INCORRECT (query logically wrong or doesn't answer the question), or UNANSWERABLE (question "
        "cannot be answered given the schema). Table 12 categorizes queries by consensus outcome: unanimous verdicts when "
        "both models agree, and 'Mixed (disagreement)' when models assign different verdicts to the same query.",
    )

    # Table 12 (semantic validation)
    headers_t12 = ["Consensus Verdict", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t12 = [
        ["CORRECT", "1,409 (65.6%)", "684 (66.2%)", "5,188 (59.9%)"],
        ["PARTIALLY_CORRECT", "43 (2.0%)", "14 (1.4%)", "259 (3.0%)"],
        ["INCORRECT", "171 (8.0%)", "94 (9.1%)", "867 (10.0%)"],
        ["Mixed (disagreement)", "522 (24.3%)", "242 (23.4%)", "2,316 (26.7%)"],
        ["UNANSWERABLE", "2 (0.1%)", "0 (0.0%)", "29 (0.3%)"],
        ["TOTAL", "2,147 (100%)", "1,034 (100%)", "8,659 (100%)"],
    ]
    add_table_with_data(
        doc,
        headers_t12,
        rows_t12,
        "Table 12. Semantic validation results via LLM consensus",
    )

    add_paragraph(
        doc,
        "Test and Dev exhibit comparable unanimous CORRECT rates: 1,409/2,147 (65.6%) and 684/1,034 (66.2%) respectively. "
        "Train shows lower correctness at 5,188/8,659 (59.9%), a 6-7 percentage point decrease that may reflect Train's "
        "larger size capturing more ambiguous cases. Unanimous INCORRECT verdicts total 1,132 queries across partitions: "
        "171/2,147 (8.0%) in Test, 94/1,034 (9.1%) in Dev, and 867/8,659 (10.0%) in Train, indicating that approximately "
        "8–10% of Spider contains queries both LLM judges independently classify as semantically wrong for their associated "
        "questions. The Mixed disagreement category comprises 522 Test queries (24.3%), 242 Dev queries (23.4%), and 2,316 "
        "Train queries (26.7%), totaling 3,080 queries (26.0% of benchmark) where frontier LLMs disagree about correctness "
        "under the fixed two-judge configuration. UNANSWERABLE verdicts are rare: 2 in Test (0.1%), 0 in Dev, and 29 in Train (0.3%), "
        "suggesting that most questions are theoretically answerable given the provided schemas. Summing the unanimous "
        "CORRECT and PARTIALLY_CORRECT categories yields an 'upper-confidence-bound' correctness estimate of approximately "
        "67.6% (Test), 67.5% (Dev), and 62.9% (Train), while the presence of substantial Mixed cases indicates inherent "
        "ambiguity in approximately one-quarter of the benchmark.",
    )

    # Table 13 (disagreement patterns)
    add_paragraph(
        doc,
        "Table 13 details the disagreement patterns within the Mixed category from Table 12, decomposing the 3,080 mixed "
        "cases into specific verdict transitions expressed as 'Gemini → GPT-5'. Percentages are computed relative to the "
        "Mixed count in each partition (denominators: 522 Test, 242 Dev, 2,316 Train), enabling comparison of disagreement "
        "pattern distributions across partitions. This breakdown reveals which semantic boundary cases dominate inter-model "
        "disagreement under the two-judge configuration.",
    )

    headers_t13 = [
        "Disagreement Pattern (Gemini → GPT-5)",
        "Spider Test",
        "Spider Dev",
        "Spider Train",
    ]
    rows_t13 = [
        ["CORRECT → PARTIALLY_CORRECT", "315 (60.3%)", "144 (59.5%)", "1,375 (59.4%)"],
        ["CORRECT → INCORRECT", "124 (23.8%)", "57 (23.6%)", "449 (19.4%)"],
        ["INCORRECT → PARTIALLY_CORRECT", "35 (6.7%)", "17 (7.0%)", "169 (7.3%)"],
        ["INCORRECT → CORRECT", "20 (3.8%)", "14 (5.8%)", "121 (5.2%)"],
        ["CORRECT → UNANSWERABLE", "14 (2.7%)", "4 (1.7%)", "94 (4.1%)"],
        ["PARTIALLY_CORRECT → CORRECT", "6 (1.1%)", "1 (0.4%)", "11 (0.5%)"],
        ["PARTIALLY_CORRECT → INCORRECT", "6 (1.1%)", "2 (0.8%)", "30 (1.3%)"],
        ["INCORRECT → UNANSWERABLE", "1 (0.2%)", "3 (1.2%)", "40 (1.7%)"],
        ["UNANSWERABLE → CORRECT", "0 (0.0%)", "0 (0.0%)", "14 (0.6%)"],
        ["UNANSWERABLE → PARTIALLY_CORRECT", "0 (0.0%)", "0 (0.0%)", "6 (0.3%)"],
        ["PARTIALLY_CORRECT → UNANSWERABLE", "0 (0.0%)", "0 (0.0%)", "5 (0.2%)"],
        ["UNANSWERABLE → INCORRECT", "1 (0.2%)", "0 (0.0%)", "2 (0.1%)"],
    ]
    add_table_with_data(
        doc,
        headers_t13,
        rows_t13,
        "Table 13. Mixed consensus breakdown (percentages relative to Mixed category)",
    )

    add_paragraph(
        doc,
        "CORRECT↔PARTIALLY_CORRECT transitions dominate disagreement across all partitions, comprising 315/522 (60.3%) "
        "of Test Mixed, 144/242 (59.5%) of Dev Mixed, and 1,375/2,316 (59.4%) of Train Mixed, totaling 1,834 queries "
        "(59.6% of all Mixed cases). This consistent ~60% majority suggests that approximately 1,800 queries occupy a "
        "semantic boundary where one model judges the SQL fully correct while the other flags minor issues (missing DISTINCT, "
        "NULL handling edge cases, date boundary conditions), representing queries that could reasonably be classified as "
        "partially correct rather than fully wrong. CORRECT↔INCORRECT disagreements constitute the second-largest pattern: "
        "124/522 (23.8%) Test, 57/242 (23.6%) Dev, 449/2,316 (19.4%) Train, totaling 630 queries (20.5% of Mixed) where "
        "models fundamentally disagree on correctness. INCORRECT↔PARTIALLY_CORRECT transitions (6.7–7.3% of Mixed) and "
        "INCORRECT↔CORRECT (3.8–5.8% of Mixed) together represent cases where at least one model judges the query semantically "
        "incorrect; these likely indicate queries with fundamental semantic errors even if judges disagree on severity, "
        "suggesting that the effective incorrect rate may exceed the unanimous INCORRECT count (8–10%) reported in Table 12. "
        "UNANSWERABLE-related transitions appear primarily in Train (CORRECT→UNANSWERABLE at 4.1%, INCORRECT→UNANSWERABLE at "
        "1.7%), while nearly absent in Test/Dev, consistent with Train's larger size sampling from rare edge cases.",
    )

    # 4.6. Performance
    doc.add_page_break()
    add_heading(doc, "4.6. Computational Performance Characteristics", level=2)
    
    add_paragraph(
        doc,
        "Table 14 presents per-query processing time by analyzer component, measured as average milliseconds per query "
        "across the partition. Times for fast analyzers (antipattern detection, syntax analysis, execution testing) reflect "
        "local computation without network latency. Schema validation times are amortized across queries (schema analysis "
        "occurs once per database and is reused for all queries sharing that database). Semantic LLM Judge times include "
        "end-to-end latency for two sequential model inferences per query (one call to Gemini 2.5 Pro, one to GPT-5) under "
        "the configured judging protocol.",
    )

    # Table 14 (performance)
    headers_t14 = ["Analyzer Component", "Test (ms/query)", "Dev (ms/query)", "Train (ms/query)"]
    rows_t14 = [
        ["Query Antipattern Detection", "1.21", "0.49", "0.56"],
        ["Query Syntax Analysis", "2.32", "0.60", "0.64"],
        ["Query Execution Testing", "3.10", "2.49", "1.19"],
        ["Schema Validation", "14.25", "16.91", "9.93"],
        ["Semantic LLM Judge", "21,271", "21,200", "20,350"],
    ]
    add_table_with_data(
        doc,
        headers_t14,
        rows_t14,
        "Table 14. Per-query processing time by analyzer component",
    )

    add_paragraph(
        doc,
        "Local/static analyzers operate within millisecond ranges: Query Antipattern Detection averages 0.49–1.21 ms/query "
        "(variation likely due to query complexity affecting AST traversal depth), Query Syntax Analysis averages 0.60–2.32 ms/query "
        "(sqlglot parsing and feature extraction), and Query Execution Testing averages 1.19–3.10 ms/query (SQLite execution "
        "under LIMIT-constrained protocol). Schema Validation, when amortized across queries, averages 9.93–16.91 ms/query "
        "(database-level introspection cost distributed over queries sharing each database). In stark contrast, the Semantic "
        "LLM Judge requires approximately 20–21 seconds per query (20,350 ms in Train, 21,200 ms in Dev, 21,271 ms in Test), "
        "representing 3–4 orders of magnitude higher latency than formal analyzers. The LLM judging time reflects sequential "
        "inference (not parallelized in this configuration) for two model calls per query, including network round-trip, "
        "API queue time, model inference, and response parsing. The consistency of LLM times across partitions (20–21 seconds) "
        "suggests that query complexity has minimal impact on LLM latency compared to baseline model inference time.",
    )

    # Table 15 (aggregate time)
    add_paragraph(
        doc,
        "Table 15 presents aggregate processing time by partition, separating runtime into fast analyzers (combined local "
        "analysis), LLM Judge (semantic validation), and I/O overhead (metric serialization, database file operations). "
        "The 'Total pipeline time' row sums these components, providing end-to-end wall-clock runtime for processing each "
        "partition under the experimental configuration. The 'LLM fraction of total' quantifies semantic validation's "
        "contribution to overall runtime.",
    )

    headers_t15 = ["Component", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t15 = [
        ["Fast analyzers (combined)", "~15 sec", "~8 sec", "~11 sec"],
        ["LLM Judge", "~12.7 hr", "~6.1 hr", "~49.0 hr"],
        ["I/O overhead", "~45 sec", "~22 sec", "~68 sec"],
        ["Total pipeline time", "~12.7 hr", "~6.1 hr", "~49.0 hr"],
        ["LLM fraction of total", "99.9%", "99.9%", "99.9%"],
    ]
    add_table_with_data(
        doc, headers_t15, rows_t15, "Table 15. Aggregate processing time by partition"
    )

    add_paragraph(
        doc,
        "Total pipeline runtime is 12.7 hours for Test (2,147 queries), 6.1 hours for Dev (1,034 queries), and 49.0 hours "
        "for Train (8,659 queries). LLM semantic validation accounts for 99.9% of runtime across all partitions, dominating "
        "total processing time by 2–3 orders of magnitude. Fast analyzers complete in seconds (approximately 15 seconds for "
        "Test, 8 seconds for Dev, 11 seconds for Train), while I/O overhead (metric serialization, DuckDB writes, JSONL "
        "output) requires under two minutes per partition (45 seconds Test, 22 seconds Dev, 68 seconds Train). The runtime "
        "scales linearly with example count (2,147 queries → 12.7 hours; 8,659 queries → 49.0 hours), reflecting the sequential "
        "query-by-query processing under the streaming architecture. Cumulative analysis time for all 11,840 examples totals "
        "approximately 67.8 hours (2.83 days) of wall-clock time under the measured configuration. The extreme LLM-judge "
        "dominance (99.9%) indicates that formal validation layers (schema/syntax/antipattern/execution) add negligible "
        "computational overhead compared to semantic validation, completing 11,840 queries in ~34 seconds total across all "
        "partitions versus 67.8 hours for LLM judging.",
    )

    # 4.7. Integrated Summary
    doc.add_page_break()
    add_heading(doc, "4.7. Integrated Quality Assessment", level=2)
    
    add_paragraph(
        doc,
        "Table 16 synthesizes key quantitative results across all validation dimensions for integrated comparative assessment. "
        "Values are drawn from Tables 1–15 using consistent units and denominators. The table structure separates results by "
        "quality dimension (DATABASE SCHEMAS, SYNTACTIC QUALITY, EXECUTABILITY, CODE QUALITY, SEMANTIC CORRECTNESS) to "
        "enable simultaneous examination of multiple validation aspects. For critical database-level issues, specific database "
        "identifiers and error descriptions are included inline to facilitate targeted remediation.",
    )

    # Table 16 (integrated assessment)
    headers_t16 = ["Quality Dimension", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t16 = [
        ["DATABASE SCHEMAS", "", "", ""],
        ["Valid schemas (%)", "90.0", "65.0", "75.3"],
        ["Warnings-only databases (count)", "0", "2", "4"],
        ["Structural FK errors (count)", "4", "7", "78"],
        [
            "    [Key structural-error DBs]",
            "book_1 (2× FK missing column), car_racing (1× FK missing column), government_shift (1× FK type mismatch)",
            "car_1 (FK type mismatch + FK target not key + 2 data violations), concert_singer (2× FK type mismatch), voter_1 (FK target not key)",
            "FK missing column (Critical): baseball_1, imdb, loan_1, restaurants, store_product; FK target not key (High): baseball_1 (20×), imdb (6×), yelp (7×); FK type mismatch (High): cre_Drama_Workshop_Groups (9×)",
        ],
        ["FK data violations (rows)", "1", "2,403", "39,523"],
        [
            "    [DBs with FK data violations]",
            "pilot_1 (1 row)",
            "flight_2 (2,400 rows), car_1 (2 rows), wta_1 (1 row)",
            "sakila_1 (38,273 rows), flight_4 (1,240 rows), hr_1 (6), college_1 (2), hospital_1 (1), allergy_1 (1)",
        ],
        ["", "", "", ""],
        ["SYNTACTIC QUALITY", "", "", ""],
        ["Parse success (%)", "100", "100", "100"],
        ["Mean complexity score", "39.5", "39.0", "40.7"],
        ["", "", "", ""],
        ["EXECUTABILITY", "", "", ""],
        ["Execution success (%)", "100", "100", "99.97"],
        ["", "", "", ""],
        ["CODE QUALITY", "", "", ""],
        ["Mean quality score (/100)", "97.7", "97.7", "98.4"],
        ["Critical antipatterns (count)", "2", "2", "13"],
        ["    [All 17 manually verified as incorrect; all 17 judged INCORRECT by both LLMs]", "", "", ""],
        ["High-severity antipatterns (%)", "14.5", "14.7", "9.5"],
        ["", "", "", ""],
        ["SEMANTIC CORRECTNESS", "", "", ""],
        ["Correct (%)", "65.6", "66.2", "59.9"],
        ["Incorrect (%)", "8.0", "9.1", "10.0"],
        ["Disputed/Mixed (%)", "24.3", "23.4", "26.7"],
    ]
    add_table_with_data(
        doc,
        headers_t16,
        rows_t16,
        "Table 16. Integrated quality assessment across all validation dimensions",
    )

    add_paragraph(
        doc,
        "The integrated assessment reveals a critical paradox: all partitions achieve perfect or near-perfect traditional "
        "metrics (100% SQL parsing success, 99.97–100% execution success, 97.7–98.4 code quality scores), yet exhibit "
        "substantial hidden defects invisible to these metrics. Database-level validation shows 10–35% of databases contain "
        "schema errors, with Dev achieving only 65.0% schema validity despite serving as the model selection partition. "
        "Semantic correctness ranges from 59.9% (Train) to 66.2% (Dev) for unanimous CORRECT verdicts, with 8–10% unanimously "
        "INCORRECT and 23.4–26.7% disputed (Mixed), indicating that approximately one-third to one-half of Spider exhibits "
        "either explicit semantic errors or inter-judge ambiguity. The 17 critical Cartesian product antipatterns (2 Test, "
        "2 Dev, 13 Train) represent confirmed incorrect queries validated through three independent methods: antipattern "
        "detection (identifying missing JOIN conditions), manual code review (confirming broken join logic), and unanimous "
        "LLM semantic judgment (both models independently classified all 17 as INCORRECT). High-severity antipattern prevalence "
        "(14.5% Test, 14.7% Dev, 9.5% Train) primarily reflects Missing GROUP BY (945 queries) and NOT IN with nullable "
        "columns (348 queries), both exploiting SQLite's permissive execution model in ways that would fail or produce wrong "
        "results in stricter database systems.",
    )

    # Add Table 17 (Critical databases requiring remediation)
    add_paragraph(
        doc,
        "Table 17 consolidates the most critical database quality issues identified, providing a prioritized remediation "
        "list for Spider maintainers.",
    )
    
    headers_t17 = ["Database ID", "Partition", "Critical Issues", "Priority"]
    rows_t17 = [
        ["sakila_1", "Train", "38,273 FK data violations + 4 empty tables", "CRITICAL"],
        ["flight_2", "Dev", "2,400 FK data violations → EVALUATION BIAS", "CRITICAL"],
        ["flight_4", "Train", "1,240 FK data violations", "CRITICAL"],
        ["hr_1", "Train", "6 FK data violations", "CRITICAL"],
        ["college_1", "Train", "2 FK data violations", "CRITICAL"],
        ["car_1", "Dev", "2 FK data violations + structural FK errors", "CRITICAL"],
        ["hospital_1", "Train", "1 FK data violation", "CRITICAL"],
        ["allergy_1", "Train", "1 FK data violation", "CRITICAL"],
        ["wta_1", "Dev", "1 FK data violation", "CRITICAL"],
        ["pilot_1", "Test", "1 FK data violation → EVALUATION BIAS (Test partition)", "CRITICAL"],
        ["book_1", "Test", "2 FK missing column (schema defect in Test partition)", "CRITICAL"],
        ["car_racing", "Test", "1 FK missing column (schema defect in Test partition)", "CRITICAL"],
        ["baseball_1", "Train", "FK missing column (Critical) + 20 FK target-not-key errors", "CRITICAL"],
        ["imdb", "Train", "FK missing column (Critical) + 6 FK target-not-key errors + all 16 tables empty", "CRITICAL"],
        ["loan_1", "Train", "FK missing column (Critical) + FK type mismatch", "CRITICAL"],
        ["store_product", "Train", "FK missing column (Critical)", "CRITICAL"],
        ["restaurants", "Train", "FK missing column (Critical) + all 3 tables empty", "CRITICAL"],
        ["cre_Drama_Workshop_Groups", "Train", "9 FK type mismatches across 18 tables", "HIGH"],
        ["academic", "Train", "All 15 tables empty + FK type mismatch", "HIGH"],
        ["yelp", "Train", "All 7 tables empty + 7 FK errors", "HIGH"],
        ["concert_singer", "Dev", "2 FK type mismatches", "MEDIUM"],
        ["voter_1", "Dev", "FK target not key error", "MEDIUM"],
    ]
    add_table_with_data(
        doc,
        headers_t17,
        rows_t17,
        "Table 17. Critical databases requiring immediate remediation prioritized by severity and impact",
    )


def generate_section5_discussion(doc):
    """
    Section 5: Discussion
    Interpretation of results, implications, and threats to validity.
    """
    print("💬 Generating Section 5: Discussion...")
    
    doc.add_page_break()
    title = doc.add_heading("5. DISCUSSION", level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True

    add_paragraph(
        doc,
        "This section interprets the empirical findings reported in Section 4, examines their implications for "
        "Text-to-SQL research and practice, and discusses the validity threats and limitations of our study.",
    )

    # 5.1. Interpretation of Findings
    add_heading(doc, "5.1. Interpretation of Key Findings", level=2)
    
    add_heading(doc, "5.1.1. The Paradox of Formal Correctness and Hidden Defects", level=3)
    add_paragraph(
        doc,
        "The results reveal a striking paradox: Spider achieves near-perfect scores on traditional quality metrics "
        "(100% parsing success, 99.97-100% execution success) while simultaneously exhibiting substantial structural "
        "and semantic defects. This finding challenges the implicit assumption in much Text-to-SQL research that "
        "successful query execution implies dataset correctness. The 89 structural foreign key errors and 41,927 "
        "referential integrity violations documented in Section 4.1 remain invisible to standard evaluation protocols "
        "because SQLite's permissive constraint enforcement allows queries to execute despite violating relational "
        "integrity [20]. This disconnect between execution success and data quality has significant implications: "
        "execution-based evaluation metrics (e.g., Execution Match) may report stable scores even when evaluating "
        "models against corrupted ground truth, leading researchers to trust benchmark results that may be systematically "
        "biased.",
    )

    add_heading(doc, "5.1.2. Partition Quality Heterogeneity and Evaluation Bias", level=3)
    add_paragraph(
        doc,
        "The quality gradient observed across partitions—Test (90.0% valid schemas), Train (75.3%), Dev (65.0%)—contradicts "
        "standard assumptions about benchmark construction. We would expect the evaluation partitions (Test/Dev) to receive "
        "more rigorous quality assurance than the training partition, yet Dev exhibits the lowest schema validity. This "
        "finding is particularly concerning because Dev serves as the primary signal for hyperparameter tuning and model "
        "selection in typical Text-to-SQL research workflows [3,7]. The presence of 2,403 referential integrity violations "
        "in Dev's flight_2 database (Section 4.1, Table 4) affects approximately 35.6% of Dev examples, potentially "
        "introducing systematic bias into validation metrics that researchers use to select between competing approaches. "
        "Models optimized against this corrupted validation signal may exhibit degraded performance on clean data while "
        "appearing superior during development, leading to suboptimal architecture choices and misleading reported improvements.",
    )

    add_heading(doc, "5.1.3. Semantic Ambiguity and the Limits of Automated Evaluation", level=3)
    add_paragraph(
        doc,
        "The semantic validation results (Section 4.5) provide the most direct evidence that formal correctness is "
        "insufficient for evaluating Text-to-SQL systems. Only 60-66% of examples achieved unanimous CORRECT verdicts "
        "from two frontier LLMs, while 8-10% were unanimously judged INCORRECT, and 23-27% fell into the Mixed disagreement "
        "category. The disagreement patterns (Table 12) reveal that Mixed cases concentrate primarily at the "
        "CORRECT↔PARTIALLY_CORRECT boundary (59-60% of Mixed), suggesting that many disputed examples involve questions "
        "that admit multiple valid SQL implementations or queries with subtle constraint differences (e.g., presence/absence "
        "of DISTINCT, treatment of NULL values, date boundary conditions). These boundary cases represent a fundamental "
        "challenge for benchmark-based evaluation: when human annotators and frontier LLMs disagree about correctness, "
        "exact-match metrics necessarily penalize plausible alternative formulations, while execution-based metrics may "
        "accept incorrect solutions that coincidentally produce matching results on small test databases.",
    )

    add_heading(doc, "5.1.4. Antipattern Prevalence and SQL Portability", level=3)
    add_paragraph(
        doc,
        "The antipattern analysis (Section 4.4) reveals that while Spider maintains high overall code quality "
        "(97.7-98.4/100), specific high-severity patterns exhibit concerning prevalence. Missing GROUP BY affects "
        "945 queries (8.0% of dataset), exploiting SQLite's permissive handling of aggregate queries that would produce "
        "errors in PostgreSQL, MySQL (with ONLY_FULL_GROUP_BY), and Oracle. Models trained on these examples learn "
        "dialect-specific SQL that may fail when deployed against stricter database systems. Similarly, the 348 queries "
        "using NOT IN with potentially nullable columns introduce subtle correctness risks under SQL's three-valued logic: "
        "when the subquery contains NULL, the predicate evaluates to UNKNOWN rather than FALSE, potentially filtering out "
        "rows unexpectedly [21]. These patterns matter less for benchmark performance (where evaluation uses the same "
        "SQLite environment that training examples target) but become critical when models are deployed in production "
        "environments with heterogeneous database systems and strict SQL conformance requirements.",
    )

    # 5.2. Implications
    doc.add_page_break()
    add_heading(doc, "5.2. Implications for Research and Practice", level=2)

    add_heading(doc, "5.2.1. Rethinking Benchmark Quality Assurance", level=3)
    add_paragraph(
        doc,
        "Our findings demonstrate that standard quality assurance practices for Text-to-SQL benchmarks—primarily manual "
        "review of SQL syntax and spot-checking query execution—are insufficient to detect the types of systematic defects "
        "documented in this study. We recommend that future benchmark construction and curation adopt multi-layer automated "
        "validation as a mandatory quality gate before release. Specifically: (i) schema validation should verify foreign "
        "key structural correctness and referential integrity on actual data, not just DDL definitions; (ii) execution testing "
        "should be supplemented with result cardinality analysis to detect Cartesian products and other pathological cases; "
        "(iii) a sample of examples (statistically sized for desired confidence intervals) should undergo LLM-based semantic "
        "validation with human adjudication of disagreement cases; (iv) all validation results should be published alongside "
        "the dataset as structured metadata, enabling researchers to filter examples by quality criteria or conduct ablation "
        "studies comparing performance on high-confidence versus disputed subsets.",
    )

    add_heading(doc, "5.2.2. Toward Quality-Aware Evaluation Protocols", level=3)
    add_paragraph(
        doc,
        "The quality heterogeneity documented in Section 4 suggests that aggregate benchmark scores may obscure important "
        "performance differences. We propose that Text-to-SQL evaluation should routinely report disaggregated metrics: "
        "(i) performance on the high-confidence subset (only examples with valid schemas and unanimous CORRECT semantic "
        "labels), providing an upper-bound estimate of model capability on clean data; (ii) performance on examples from "
        "databases with structural defects versus clean databases, quantifying sensitivity to schema quality; (iii) performance "
        "stratified by semantic verdict category (correct/incorrect/disputed), revealing whether models are preferentially "
        "learning from high-quality or low-quality training examples; (iv) separate reporting for examples with and without "
        "high-severity antipatterns, indicating whether models learn dialect-specific versus portable SQL. This disaggregated "
        "evaluation would enable more nuanced model comparison and help the community understand which quality dimensions most "
        "strongly influence Text-to-SQL performance.",
    )

    add_heading(doc, "5.2.3. Remediation Priorities for Spider", level=3)
    add_paragraph(
        doc,
        "Based on impact analysis, we recommend the following remediation priorities for Spider 1.0: (i) CRITICAL: repair "
        "Dev's flight_2 database (2,403 violations affecting 35.6% of validation examples) before any further benchmark use "
        "in model selection or published evaluations, as these violations directly corrupt validation metrics; (ii) HIGH: "
        "address Train's catastrophic databases (sakila_1 with 38,273 violations; flight_4 with 1,240 violations) to reduce "
        "training noise and prevent models from learning to accommodate referential integrity violations; (iii) MEDIUM: fix "
        "structural foreign key errors in 19 databases exhibiting FK Type Mismatch or FK Target Not Key errors, as these "
        "represent fundamental schema design flaws that compromise JOIN operation semantics; (iv) LOW: resolve or document "
        "the seven databases with 0% data coverage, either by restoring missing data or explicitly labeling them as schema-only "
        "examples unsuitable for execution-based evaluation. Beyond these specific fixes, we recommend that Spider maintainers "
        "publish an official 'high-confidence' subset excluding all examples with schema defects or semantic disputes, enabling "
        "researchers to conduct controlled experiments on known-clean data.",
    )

    add_heading(doc, "5.2.4. Implications for Model Development", level=3)
    add_paragraph(
        doc,
        "The presence of 867-1,132 explicitly incorrect examples (Section 4.5) and hundreds of high-severity antipatterns "
        "(Section 4.4) raises important questions about fine-tuning strategies. Standard supervised learning on the full "
        "Spider dataset necessarily trains models on incorrect examples, potentially degrading performance on correct examples "
        "through negative transfer. We recommend that practitioners consider quality-aware training strategies: (i) filter "
        "training data to exclude unanimous INCORRECT examples before fine-tuning; (ii) implement example weighting schemes "
        "that down-weight disputed (Mixed) examples while emphasizing high-confidence correct examples; (iii) incorporate "
        "schema validation into the training loop, masking or skipping examples from databases with foreign key errors; "
        "(iv) augment training objectives with antipattern penalties, explicitly encouraging models to generate standard-compliant "
        "SQL rather than dialect-specific constructs. These strategies remain unexplored in existing Text-to-SQL literature but "
        "may yield meaningful improvements by focusing model learning on the highest-quality subset of training data.",
    )

    # 5.3. Threats to Validity
    doc.add_page_break()
    add_heading(doc, "5.3. Threats to Validity", level=2)

    add_heading(doc, "5.3.1. Internal Validity", level=3)
    add_paragraph(
        doc,
        "LLM-based semantic validation introduces potential systematic biases that may affect result interpretation. Both "
        "Gemini 2.5 Pro and GPT-5 were likely exposed to Spider during pre-training, potentially creating evaluation "
        "leakage where models recognize benchmark examples and produce judgments based on memorized annotations rather than "
        "independent reasoning. This threat is partially mitigated by the substantial disagreement rate (23-27% Mixed verdicts) "
        "and the high frequency of INCORRECT judgments (8-10%), suggesting that models are not simply reproducing training "
        "labels, but cannot be fully eliminated without using models demonstrably untrained on Spider. Additionally, the "
        "two-model consensus approach, while providing higher confidence than single-model judgments, remains vulnerable to "
        "shared failure modes: both models may systematically misunderstand specific SQL constructs (e.g., NULL handling in "
        "subqueries, duplicate elimination semantics) or database domains (e.g., sports statistics, academic records), leading "
        "to unanimous incorrect verdicts that we classify as ground truth. The provision of sampled cell values (Section 2.8) "
        "introduces sampling bias: our samples may not capture edge cases (e.g., extreme values, rare categories, NULL prevalence) "
        "that are necessary for accurate semantic judgment, particularly for queries involving statistical aggregates or "
        "distribution-dependent conditions.",
    )

    add_heading(doc, "5.3.2. External Validity", level=3)
    add_paragraph(
        doc,
        "This study validates the framework exclusively on Spider 1.0, a cross-domain benchmark with academic-style questions "
        "and small databases. Generalization to other Text-to-SQL benchmarks with different characteristics remains unproven. "
        "BIRD [12] employs substantially larger databases (up to 200 tables, gigabytes of data) where referential integrity "
        "violations may have different prevalence and impact; WikiSQL [11] uses single-table queries exclusively, rendering "
        "foreign key validation and JOIN analysis inapplicable; synthetic datasets generated by recent LLMs [13,14] may exhibit "
        "entirely different error distributions reflecting model artifacts rather than human annotation patterns. The validation "
        "framework's computational cost (67.8 hours for 11,840 examples) may become prohibitive for much larger benchmarks "
        "(e.g., OmniSQL with 2.5M examples [14]), requiring sampling strategies that introduce additional validity threats. "
        "The antipattern catalog (Section 2.6) reflects best practices for production SQL but may not align with conventions "
        "in specific benchmark construction methodologies, potentially flagging intentional design choices as defects.",
    )

    add_heading(doc, "5.3.3. Construct Validity", level=3)
    add_paragraph(
        doc,
        "Our operationalization of 'dataset quality' through five dimensions (schema integrity, syntactic structure, executability, "
        "antipatterns, semantic correctness) necessarily omits other potentially important quality aspects. We do not assess "
        "question naturalness or linguistic diversity, which may influence model performance independently of SQL correctness; "
        "we do not evaluate whether database schemas accurately represent their claimed domains (e.g., whether the 'concert_singer' "
        "database schema matches real-world concert management systems); we do not verify that natural language questions reflect "
        "realistic information needs or query patterns from actual database users [1]. The complexity scoring function (Section 2.4) "
        "employs equal weights for different SQL features, but certain constructs (e.g., correlated subqueries, HAVING clauses) may "
        "contribute disproportionately to model difficulty in ways our metric does not capture. The semantic correctness categories "
        "(CORRECT, PARTIALLY_CORRECT, INCORRECT, UNANSWERABLE, Mixed) provide only coarse-grained distinctions; finer-grained error "
        "taxonomies (e.g., distinguishing missing JOIN conditions from incorrect aggregation functions from wrong filtering predicates) "
        "might reveal more actionable patterns but would require substantially more complex annotation protocols.",
    )

    add_heading(doc, "5.3.4. Reliability", level=3)
    add_paragraph(
        doc,
        "The LLM-based semantic validation exhibits inherent non-determinism even when prompts and generation settings are held constant, "
        "as frontier models can demonstrate residual variability and prompt sensitivity in practice. "
        "We did not conduct repeated evaluations to quantify inter-run reliability; therefore, the specific correctness percentages "
        "(e.g., 65.6% CORRECT for Test) should be interpreted as point estimates rather than stable measurements. Systematic "
        "inter-annotator reliability studies comparing LLM verdicts with expert human judgments on a statistically powered sample "
        "would be necessary to calibrate the false positive and false negative rates of our automated semantic validation. The "
        "schema validation component's reliability depends on correct implementation of foreign key constraint semantics across "
        "SQL dialects; while we extensively tested our validation logic (Section 2.3), subtle dialect differences in constraint "
        "interpretation may cause our validator to flag valid constructs as errors or miss actual violations in edge cases.",
    )


def generate_section6_conclusion(doc):
    """
    Section 6: Conclusion
    Summary, contributions, future work, and final remarks.
    """
    print("✅ Generating Section 6: Conclusion...")
    
    doc.add_page_break()
    title = doc.add_heading("6. CONCLUSION", level=1)
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True

    # 6.1. Summary
    add_heading(doc, "6.1. Summary of Contributions", level=2)
    
    add_paragraph(
        doc,
        "This paper introduces the first comprehensive automated framework for multi-dimensional quality assessment of "
        "Text-to-SQL datasets and demonstrates its application to the widely-used Spider 1.0 benchmark. Our validation "
        "system integrates five complementary quality dimensions—database schema integrity, SQL syntactic structure, "
        "dynamic executability, antipattern detection, and LLM-based semantic verification—into a unified pipeline that "
        "processes datasets at scale while producing actionable quality metrics and detailed error reports.",
    )

    add_paragraph(
        doc,
        "The empirical analysis of all 11,840 Spider examples across three partitions (Test/Dev/Train) reveals a critical "
        "finding: benchmarks can achieve near-perfect traditional correctness metrics (100% parsing, 99.97% execution success) "
        "while harboring substantial hidden defects that directly impact model training and evaluation validity. We document "
        "89 structural foreign key errors, 41,927 referential integrity violations concentrated in two catastrophic databases, "
        "71 empty tables, and approximately 1,100 semantically incorrect query annotations. Most significantly, we identify "
        "that Dev—the partition used for model selection in standard workflows—exhibits the lowest schema quality (65% valid "
        "databases) and contains 2,403 referential integrity violations affecting 35.6% of validation examples, introducing "
        "systematic bias into the hyperparameter tuning process.",
    )

    add_paragraph(
        doc,
        "The semantic validation results demonstrate that only 60-66% of Spider examples achieve unanimous correctness "
        "judgments from frontier LLMs, with 8-10% unanimously incorrect and 23-27% disputed. This finding challenges the "
        "standard practice of treating benchmark labels as ground truth and suggests that Text-to-SQL evaluation must "
        "account for inherent ambiguity in question-SQL alignment. The antipattern analysis reveals that while Spider "
        "maintains high overall code quality (97.7-98.4/100), specific high-severity patterns—particularly Missing GROUP BY "
        "(945 occurrences) and NOT IN with nullable columns (348 occurrences)—introduce portability and correctness risks "
        "that become critical when deploying models in production environments with stricter SQL conformance requirements.",
    )

    add_paragraph(
        doc,
        "We release the validation framework as open-source software, enabling reproducible quality assessment of existing "
        "benchmarks and providing benchmark creators with automated quality assurance tooling. The framework is generic by "
        "design: its modular architecture supports application to any Text-to-SQL dataset conforming to standard formats, "
        "requiring only configuration changes rather than code modifications to validate alternative benchmarks.",
    )

    # 6.2. Implications
    add_heading(doc, "6.2. Broader Implications", level=2)
    
    add_paragraph(
        doc,
        "Our findings suggest that the Text-to-SQL research community should reconsider several standard practices. First, "
        "benchmark quality should be treated as a first-class research concern requiring systematic measurement and transparent "
        "reporting, not an implicit assumption. Just as machine learning research has embraced dataset documentation frameworks "
        "like Datasheets [23] and Data Statements [24], Text-to-SQL benchmarks should include structured quality metadata "
        "enabling researchers to filter examples by validation criteria and conduct quality-aware ablation studies. Second, "
        "evaluation protocols should routinely disaggregate performance by quality dimensions (schema validity, semantic confidence, "
        "antipattern presence) to reveal whether apparent model improvements stem from genuine capability advances or from "
        "overfitting to dataset defects. Third, model training strategies should incorporate quality-aware filtering or weighting "
        "to prevent learning from incorrect examples—a straightforward intervention made possible by automated validation but "
        "largely unexplored in existing Text-to-SQL literature.",
    )

    add_paragraph(
        doc,
        "Beyond the immediate Text-to-SQL domain, this work contributes to broader conversations about benchmark reliability "
        "in natural language processing and program synthesis. Recent studies have documented annotation errors and ambiguity "
        "in widely-used NLP benchmarks [25,26], but most focus on inter-annotator agreement or manual re-annotation studies "
        "that cannot scale to modern dataset sizes. Our automated validation approach demonstrates that for tasks with formal "
        "output specifications (SQL, code, logical forms), multi-layer programmatic validation can expose systematic quality "
        "issues invisible to execution-based metrics, providing a template for quality assessment in other structured prediction "
        "domains.",
    )

    # 6.3. Limitations and Future Work
    add_heading(doc, "6.3. Limitations and Future Work", level=2)
    
    add_paragraph(
        doc,
        "Several limitations of the current study suggest directions for future research. First, our LLM-based semantic "
        "validation requires approximately 67.8 hours for 11,840 examples due to sequential API calls, making full-dataset "
        "validation expensive for iterative benchmark development. Future work should investigate more efficient validation "
        "strategies: training specialized semantic equivalence classifiers on LLM-labeled data, developing active learning "
        "approaches that prioritize uncertain examples for expensive validation, or implementing dynamic parallelism that "
        "adapts to API rate limits to maximize throughput. Second, we validate only Spider 1.0; systematic comparison across "
        "multiple benchmarks (BIRD, WikiSQL, CHASE-SQL, synthetic datasets) would reveal whether the quality patterns we observe "
        "represent Spider-specific issues or reflect broader systematic challenges in Text-to-SQL benchmark construction. We "
        "plan to extend the framework to these datasets in subsequent work.",
    )

    add_paragraph(
        doc,
        "Third, our semantic validation employs only two LLM judges; expanding the committee to include additional frontier "
        "models or incorporating specialized SQL evaluation models might improve consensus reliability. However, this must be "
        "balanced against increased computational cost and the risk that additional models trained on similar data will exhibit "
        "correlated failure modes without improving judgment quality. Fourth, we identify quality issues but do not automatically "
        "repair them; future research should investigate whether defects can be corrected programmatically (e.g., foreign key "
        "type mismatches via automatic type casting, missing GROUP BY via deterministic SQL transformations) or require human "
        "intervention, and develop confidence calibration methods to distinguish high-confidence automated fixes from cases "
        "requiring expert review.",
    )

    add_paragraph(
        doc,
        "Fifth, we do not conduct controlled experiments measuring the impact of quality defects on model performance—such "
        "experiments would require training models on original versus quality-filtered versions of Spider and comparing their "
        "performance on clean hold-out data. This represents valuable future work that could quantify the practical benefits "
        "of quality-aware training and evaluation. Sixth, our antipattern catalog reflects general SQL best practices but may "
        "not capture Text-to-SQL-specific quality dimensions; future work should investigate whether certain 'antipatterns' "
        "actually represent intentional benchmark design choices (e.g., using SELECT * to test column selection capabilities) "
        "and develop task-specific quality criteria beyond generic SQL coding standards.",
    )

    add_paragraph(
        doc,
        "Finally, we focus exclusively on English-language benchmarks with natural questions and do not address quality "
        "assessment for multilingual Text-to-SQL datasets, domain-specific SQL dialects (e.g., Spark SQL, BigQuery SQL), "
        "or conversational multi-turn Text-to-SQL tasks. Extending the framework to these settings would require additional "
        "validation dimensions (e.g., cross-lingual semantic consistency, dialect-specific feature coverage, dialogue coherence "
        "metrics) that remain open research challenges.",
    )

    # 6.4. Final Remarks
    add_heading(doc, "6.4. Final Remarks", level=2)
    
    add_paragraph(
        doc,
        "The Text-to-SQL field has achieved remarkable progress over the past five years, with frontier LLM-based systems "
        "approaching human-level performance on established benchmarks. This progress creates an imperative to ensure that "
        "the benchmarks guiding development and measuring capability are themselves of high quality and free from systematic "
        "defects that could invalidate research conclusions. Our finding that a widely-trusted benchmark can simultaneously "
        "exhibit high execution success and severe hidden quality issues serves as a cautionary tale: standard evaluation "
        "metrics may provide false confidence in benchmark reliability.",
    )

    add_paragraph(
        doc,
        "We advocate for a cultural shift in Text-to-SQL research toward treating benchmark quality as measurable, reportable, "
        "and improvable. The validation framework and empirical findings presented in this paper provide both the technical "
        "infrastructure and the empirical motivation for this shift. By making comprehensive dataset validation accessible "
        "through open-source tooling and demonstrating that automated quality assessment can scale to real-world benchmarks, "
        "we hope to enable researchers to build on higher-quality data foundations and to make benchmark evaluation more "
        "trustworthy and scientifically rigorous.",
    )

    add_paragraph(
        doc,
        "The Spider benchmark has been instrumental in driving Text-to-SQL progress, and our critique should be understood "
        "as a constructive contribution aimed at strengthening rather than undermining this valuable community resource. We "
        "encourage Spider maintainers and users to adopt the validation framework for quality monitoring, to incorporate "
        "quality metadata into future releases, and to consider publishing curated high-confidence subsets alongside the full "
        "dataset. For the broader research community, we hope this work establishes multi-layer automated validation as a "
        "standard practice in benchmark development and demonstrates that investing in dataset quality assurance yields "
        "scientific dividends through more reliable model evaluation and more trustworthy research conclusions.",
    )


def generate_all_sections():
    """Generate all sections in a single DOCX document."""
    print("🚀 Generating Sections 4, 5, and 6 for e-Informatica Journal...")
    print("=" * 80)
    
    # Create document
    doc = Document()

    # Configure styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Generate each section
    generate_section4_results(doc)
    generate_section5_discussion(doc)
    generate_section6_conclusion(doc)

    # Save document
    print("\n" + "=" * 80)
    print(f"💾 Saving document: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print(f"✅ Document saved successfully!")
    
    print("\n" + "=" * 80)
    print("📄 DOCUMENT STRUCTURE SUMMARY")
    print("=" * 80)
    print("\n✅ SECTION 4: RESULTS (pure empirical findings)")
    print("   4.1. Database Schema Integrity")
    print("       • Table 1: Schema validation overview")
    print("       • Table 2: FK error distribution by type")
    print("       • Table 3: FK errors by database ID (detailed)")
    print("       • Table 4: FK data violations by database")
    print("       • Table 5: Empty tables by database")
    print("   4.2. SQL Syntactic Structure and Complexity")
    print("       • Table 6: Syntactic analysis")
    print("       • Table 7: Difficulty distribution")
    print("       • Table 8: JOIN distribution")
    print("   4.3. Dynamic Query Executability")
    print("       • Table 9: Execution results")
    print("   4.4. SQL Code Quality and Antipatterns")
    print("       • Table 10: Quality metrics")
    print("       • Table 11: Antipattern distribution")
    print("           ✅ CRITICAL: All 17 Cartesian products manually verified as incorrect")
    print("           ✅ All 17 also judged INCORRECT by both LLMs independently")
    print("   4.5. Semantic Correctness Assessment")
    print("       • Table 12: LLM consensus results")
    print("       • Table 13: Disagreement patterns (with interpretation)")
    print("           ✅ CORRECT↔PARTIALLY_CORRECT: Can be considered partially correct")
    print("           ✅ INCORRECT↔PARTIALLY_CORRECT + INCORRECT↔UNANSWERABLE: Should be considered incorrect")
    print("   4.6. Computational Performance")
    print("       • Table 14: Per-query processing time")
    print("       • Table 15: Aggregate processing time")
    print("   4.7. Integrated Quality Assessment")
    print("       • Table 16: Cross-dimensional synthesis with Cartesian product note")
    print("       • Table 17: Critical databases requiring remediation")
    
    print("\n✅ SECTION 5: DISCUSSION (interpretation & implications)")
    print("   5.1. Interpretation of Key Findings")
    print("       5.1.1. Paradox of formal correctness and hidden defects")
    print("       5.1.2. Partition quality heterogeneity and evaluation bias")
    print("       5.1.3. Semantic ambiguity and limits of automated evaluation")
    print("       5.1.4. Antipattern prevalence and SQL portability")
    print("   5.2. Implications for Research and Practice")
    print("       5.2.1. Rethinking benchmark quality assurance")
    print("       5.2.2. Toward quality-aware evaluation protocols")
    print("       5.2.3. Remediation priorities for Spider")
    print("       5.2.4. Implications for model development")
    print("   5.3. Threats to Validity")
    print("       5.3.1. Internal validity (LLM bias, consensus limitations)")
    print("       5.3.2. External validity (generalization to other benchmarks)")
    print("       5.3.3. Construct validity (quality operationalization)")
    print("       5.3.4. Reliability (measurement stability)")
    
    print("\n✅ SECTION 6: CONCLUSION (summary & future work)")
    print("   6.1. Summary of Contributions")
    print("   6.2. Broader Implications")
    print("   6.3. Limitations and Future Work")
    print("   6.4. Final Remarks")
    
    print("\n" + "=" * 80)
    print("📊 DOCUMENT STATISTICS")
    print("=" * 80)
    print(f"   Total Tables: 17 (ALL tables from original v2 included)")
    print(f"   Total Subsections: 4.1-4.7 (7), 5.1-5.3 (7), 6.1-6.4 (4)")
    print(f"   Academic Style: e-Informatica compliant")
    print(f"   Results-only Section 4: No interpretation or discussion")
    print(f"   Comprehensive Discussion: All implications in Section 5")
    print(f"   Threats to Validity: Complete 4-subsection treatment")
    print(f"   Citation Format: Numeric references [1-26]")
    
    print("\n" + "=" * 80)
    print("🎯 KEY IMPROVEMENTS OVER PREVIOUS VERSION")
    print("=" * 80)
    print("   ✅ ALL 17 tables from original included (none missing)")
    print("   ✅ Table 3: FK errors by database ID (detailed)")
    print("   ✅ Table 13: ALL Mixed disagreement patterns (12 combinations from actual data)")
    print("   ✅ Table 17: Critical databases - ALL with FK violations marked CRITICAL")
    print("   ✅ Cartesian products: Manual verification + LLM confirmation noted")
    print("   ✅ Mixed interpretation: CORRECT↔PARTIALLY vs INCORRECT patterns explained")
    print("   ✅ Integrated table (16): Includes critical antipattern count + note")
    print("   ✅ Clear separation: Results (4) vs Discussion (5) vs Conclusion (6)")
    print("   ✅ No methodology redundancy (references Section 2 only)")
    print("   ✅ Comprehensive Threats to Validity (4 subsections)")
    
    print("\n" + "=" * 80)
    print("🔬 CRITICAL FINDINGS PROPERLY DOCUMENTED")
    print("=" * 80)
    print("   ✅ 17 Cartesian products: All manually verified as incorrect queries")
    print("   ✅ Independent confirmation: All 17 judged INCORRECT by both LLMs")
    print("   ✅ FK Data Violations: ALL databases marked CRITICAL (including Test with 1)")
    print("   ✅ Mixed category - ALL 12 patterns from data:")
    print("       • CORRECT↔PARTIALLY_CORRECT (59-60%): Partially correct")
    print("       • CORRECT↔INCORRECT (19-24%): Major disagreement")
    print("       • INCORRECT↔PARTIALLY_CORRECT (6.7-7.3%): Lean incorrect")
    print("       • INCORRECT↔CORRECT (3.8-5.8%): Major disagreement")
    print("       • CORRECT↔UNANSWERABLE (1.7-4.1%): Ambiguous")
    print("       • PARTIALLY_CORRECT↔CORRECT (0.4-1.1%): Minor issue")
    print("       • PARTIALLY_CORRECT↔INCORRECT (0.8-1.3%): Lean incorrect")
    print("       • INCORRECT↔UNANSWERABLE (0.2-1.7%): Should be incorrect")
    print("       • Plus 4 UNANSWERABLE-related patterns (Train only)")
    print("   ✅ Integrated table shows Critical antipattern count with validation note")
    print("   ✅ Table 17 provides actionable remediation priorities")
    
    print("\n✨ Generation complete! All tables included, all critical findings documented.")


if __name__ == "__main__":
    generate_all_sections()

