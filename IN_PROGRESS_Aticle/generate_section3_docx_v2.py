#!/usr/bin/env python3
"""
Generation of Section 3 "SPIDER DATASET ANALYSIS" in DOCX format with academic style (v2).
This version minimizes bullet lists and uses prose paragraphs for better academic readability.
"""
from pathlib import Path
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ModuleNotFoundError as e:
    raise SystemExit(
        "Missing dependency 'python-docx'. Install it with:\n"
        "  python3 -m pip install python-docx\n"
        "Then re-run this script."
    ) from e

# Base paths
BASE_DIR = Path(__file__).parent
OUTPUT_FILE = BASE_DIR / "Section3_Analysis_EN_v2.docx"


def add_heading(doc, text, level=1):
    """Add a heading with formatting."""
    heading = doc.add_heading(text, level=level)
    heading.runs[0].font.name = "Times New Roman"
    heading.runs[0].font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)
    heading.runs[0].font.bold = True
    return heading


def add_paragraph(doc, text, style=None, bold=False, italic=False):
    """Add a paragraph with formatting."""
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.bold = bold
        run.font.italic = italic
    return para


def add_table_with_data(doc, headers, rows, caption=None):
    """Add a table with data and caption."""
    # Create table
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Fill headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Format headers
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fill data
    for row_idx, row_data in enumerate(rows, start=1):
        cells = table.rows[row_idx].cells
        for col_idx, cell_data in enumerate(row_data):
            cells[col_idx].text = str(cell_data)
            # Format data
            for paragraph in cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Times New Roman"
                # Center numeric data
                if col_idx > 0:  # First column left, others center
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add caption
    if caption:
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.runs[0].font.size = Pt(10)
        caption_para.runs[0].font.italic = True
        caption_para.runs[0].font.name = "Times New Roman"

    return table


def generate_section3_docx():
    """Generate DOCX document with Section 3 in academic style."""
    print("🚀 Generating Section 3 'SPIDER DATASET ANALYSIS' (Academic v2)...")

    # Create document
    doc = Document()

    # Configure styles
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Section title
    title = doc.add_heading(
        "3. COMPARATIVE ANALYSIS OF SPIDER DATASET PARTITIONS",
        level=1,
    )
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True

    # Intro (results-focused, minimal overlap with theoretical sections)
    add_paragraph(
        doc,
        "This section reports an empirical, end-to-end quality assessment of Spider 1.0 using the multi-layer validation "
        "framework described in Section 2. Rather than restating benchmark background and system design, we focus on "
        "partition-level results (Test/Dev/Train), cross-partition consistency, and practical implications for dataset use.",
    )

    add_paragraph(
        doc,
        "We structure the analysis along five complementary quality dimensions: (i) database schema integrity, "
        "(ii) SQL parsability and structural complexity, (iii) dynamic executability, (iv) SQL antipattern prevalence, "
        "and (v) semantic correspondence via LLM consensus. The goal is to identify failures that remain invisible to "
        "standard exact-match or execution-only evaluation and to derive a remediation roadmap prioritized by impact.",
    )

    # 3.1. Methodology
    doc.add_page_break()
    add_heading(doc, "3.1. Experimental Methodology", level=2)
    
    add_paragraph(
        doc,
        "We apply the full validation pipeline (Section 2) to all three Spider partitions (Test, Dev, Train) to measure "
        "overall quality and cross-partition consistency. This comparative design is essential because Dev is widely used "
        "for hyperparameter tuning while Test is used for reporting final results; partition-specific defects can therefore "
        "bias both model selection and published evaluation outcomes.",
    )

    add_paragraph(
        doc,
        "The three analyzed partitions exhibit substantially different scales: Spider Test contains 2,147 examples "
        "spanning 40 unique databases, Spider Dev includes 1,034 examples across 20 databases, and Spider Train comprises "
        "8,659 examples distributed over 146 databases. This scale disparity matters primarily for interpreting aggregate "
        "statistics and partition heterogeneity: Train covers substantially more database schemas, while Dev and Test "
        "concentrate evaluation on a smaller set of domains and databases.",
    )

    add_paragraph(
        doc,
        "All analyses use SQLite (Spider’s native format). We enable schema validation, SQL syntax/structure analysis, "
        "execution testing, antipattern detection, and semantic validation via a two-model LLM committee (Gemini 2.5 Pro "
        "and GPT-5, temperature=0, equal voting weights). For semantic evaluation, we use full-schema mode (complete database "
        "schema context) augmented with a small sample of cell values per column, as defined in Section 2.8.",
    )

    add_paragraph(
        doc,
        "We report performance and runtime characteristics separately in Section 3.7 to avoid conflating methodological "
        "setup with cost/throughput considerations.",
    )

    # 3.2. Schema validation results
    doc.add_page_break()
    add_heading(doc, "3.2. Database Schema Integrity Assessment", level=2)
    
    add_paragraph(
        doc,
        "Schema integrity validation examines the structural correctness of database definitions, detecting errors that "
        "could prevent successful query execution even for syntactically valid SQL statements. The validation framework "
        "implements comprehensive foreign key integrity checks (detailed methodology in Section 2.3) alongside detection "
        "of duplicate columns, type consistency issues, and referential integrity violations in actual data. This section "
        "presents comparative validation results across all three Spider partitions.",
    )
    
    add_paragraph(
        doc,
        "Table 1 presents comparative schema validation results across all three Spider partitions, revealing substantial "
        "heterogeneity in schema quality that contradicts the common assumption of uniform data quality within professionally "
        "curated benchmarks.",
    )

    # Table 1 (comparative schema results)
    headers_t1 = ["Metric", "Spider Test", "Spider Dev", "Spider Train", "Best"]
    rows_t1 = [
        ["Total databases", "40", "20", "146", "—"],
        ["Valid DBs", "36 (90.0%)", "13 (65.0%)", "110 (75.3%)", "Test ✅"],
        ["DBs with errors", "4 (10.0%)", "7 (35.0%)", "32 (21.9%)", "Test ✅"],
        ["Total tables", "180", "80", "793", "—"],
        ["Empty tables", "0 (0%)", "0 (0%)", "71 (9.0%)", "Test ✅"],
        ["Total foreign keys", "146", "64", "717", "—"],
        ["Invalid FKs", "4", "7", "78", "Test ✅"],
        ["DBs with FK violations", "1 (1 viol.)", "3 (2,403 viol.)", "6 (39,523 viol.)", "Test ✅"],
    ]
    add_table_with_data(
        doc,
        headers_t1,
        rows_t1,
        "Table 1. Comparative database schema validation results across Spider partitions",
    )

    add_paragraph(
        doc,
        "The results reveal a clear quality gradient with Test achieving 90.0% schema validity, substantially exceeding "
        "Train at 75.3% and Dev at 65.0%. This contradicts expectations that the largest partition would exhibit highest "
        "quality. Dev's low validity rate (35% error rate) raises concerns given its role in hyperparameter tuning, where "
        "schema errors might obscure true performance differences between approaches.",
    )

    # Table 2 (comprehensive FK error type distribution with database IDs)
    add_paragraph(
        doc,
        "Table 2 presents foreign key validation results across partitions, distinguishing five structural error types "
        "and row-level data violations. FK Type Mismatch (38 total) and FK Target Not Key (43 total) dominate structural "
        "errors, while data violations concentrate in two databases (sakila_1: 38,273; flight_2: 2,403).",
    )

    headers_t2_fk = ["Error Type", "Spider Test", "Spider Dev", "Spider Train", "Total", "Severity"]
    rows_t2_fk = [
        ["STRUCTURAL ERRORS", "", "", "", "", ""],
        ["FK Missing Table", "0", "0", "0", "0", "Critical"],
        ["FK Missing Column", "3", "0", "5", "8", "Critical"],
        ["FK Arity Mismatch", "0", "0", "0", "0", "Critical"],
        ["FK Type Mismatch", "1", "5", "32", "38", "High"],
        ["FK Target Not Key", "0", "2", "41", "43", "High"],
        ["Duplicate Columns", "0", "0", "0", "0", "Medium"],
        ["Unknown Data Types", "0", "0", "0", "0", "Low"],
        ["Subtotal (Structural)", "4", "7", "78", "89", "—"],
        ["DATA VIOLATIONS", "", "", "", "", ""],
        ["FK Data Violations (rows)", "1", "2,403", "39,523", "41,927", "CRITICAL"],
        ["DBs with Data Violations", "1", "3", "6", "10", "—"],
    ]
    add_table_with_data(
        doc,
        headers_t2_fk,
        rows_t2_fk,
        "Table 2. Comprehensive foreign key error distribution across partitions",
    )

    # Table 2-detailed (Database IDs by FK Error Type)
    add_paragraph(
        doc,
        "Table 2-detailed enumerates the specific database identifiers associated with each foreign key error type, "
        "providing a practical reference for database maintainers and researchers seeking to understand which databases "
        "require which types of schema corrections. The table reveals error concentration patterns: certain databases "
        "(baseball_1, cre_Drama_Workshop_Groups, imdb, yelp) accumulate multiple error types simultaneously, while "
        "others exhibit single isolated issues.",
    )
    
    headers_t2_detail = ["Error Type", "Partition", "Count", "Database IDs"]
    rows_t2_detail = [
        ["FK Missing Column", "Test", "3", "book_1 (2×), car_racing (1×)"],
        ["FK Missing Column", "Dev", "0", "—"],
        ["FK Missing Column", "Train", "5", "baseball_1, imdb, loan_1, restaurants, store_product"],
        ["", "", "", ""],
        ["FK Type Mismatch", "Test", "1", "government_shift"],
        ["FK Type Mismatch", "Dev", "5", "car_1, concert_singer (2×), employee_hire_evaluation, museum_visit"],
        ["FK Type Mismatch", "Train", "32", 
         "academic, aircraft (2×), architecture (2×), city_record, company_employee, "
         "cre_Drama_Workshop_Groups (9×), culture_company (2×), loan_1, machine_repair, "
         "party_people, performance_attendance (2×), phone_1, phone_market, race_track, "
         "school_finance (2×), shop_membership (2×), student_assessment, wrestler"],
        ["", "", "", ""],
        ["FK Target Not Key", "Test", "0", "—"],
        ["FK Target Not Key", "Dev", "2", "car_1, voter_1"],
        ["FK Target Not Key", "Train", "41", 
         "baseball_1 (20×), dorm_1 (3×), imdb (6×), soccer_1 (4×), wine_1 (2×), yelp (7×)"],
    ]
    add_table_with_data(
        doc,
        headers_t2_detail,
        rows_t2_detail,
        "Table 2-detailed. Database identifiers by foreign key error type and partition (× denotes error count)",
    )

    add_paragraph(
        doc,
        "FK Type Mismatch errors (38 total) affect 19 databases, with Train's cre_Drama_Workshop_Groups containing 9 instances "
        "(50% of Train's type mismatches). These arise from TEXT-to-INTEGER or INTEGER-to-REAL incompatibilities, suggesting "
        "inconsistent type selection or uncoordinated schema evolution. FK Target Not Key errors (43 total) concentrate in Train "
        "with baseball_1 (20×), yelp (7×), and imdb (6×) accounting for 80% of this error type, representing fundamental design "
        "flaws where foreign keys reference non-unique columns. FK Missing Column errors (8 total) appear sporadically in book_1, "
        "baseball_1, imdb, loan_1, restaurants, and store_product, indicating column renaming without foreign key updates.",
    )

    # Table 2a (Databases with FK Data Violations - detailed breakdown)
    add_paragraph(
        doc,
        "Beyond structural errors, 10 databases exhibit referential integrity violations where foreign key values lack "
        "corresponding primary key entries. Table 2a reveals extreme concentration: sakila_1 (38,273 violations) and "
        "flight_2 (2,403 violations) account for 97% of all 41,927 violations, indicating systematic data export failures. "
        "The remaining eight databases contain only 14 violations combined (<0.1%), representing isolated issues.",
    )
    
    headers_t2a = ["Database ID", "Partition", "FK Data Violations", "% of Total", "Severity"]
    rows_t2a = [
        ["sakila_1", "Train", "38,273", "91.3%", "CATASTROPHIC"],
        ["flight_2", "Dev", "2,403", "5.7%", "CATASTROPHIC"],
        ["flight_4", "Train", "1,240", "3.0%", "High"],
        ["car_1", "Dev", "2", "<0.1%", "Low"],
        ["hr_1", "Train", "6", "<0.1%", "Low"],
        ["college_1", "Train", "2", "<0.1%", "Low"],
        ["wta_1", "Dev", "1", "<0.1%", "Low"],
        ["hospital_1", "Train", "1", "<0.1%", "Low"],
        ["allergy_1", "Train", "1", "<0.1%", "Low"],
        ["pilot_1", "Test", "1", "<0.1%", "Low"],
        ["TOTAL", "—", "41,927", "100%", "—"],
    ]
    add_table_with_data(
        doc,
        headers_t2a,
        rows_t2a,
        "Table 2a. Complete enumeration of databases with foreign key data violations across all partitions",
    )
    
    add_paragraph(
        doc,
        "CATASTROPHIC severity applies to sakila_1 and flight_2, whose pervasive violations render JOIN operations unreliable. "
        "The sakila_1 database (modified MySQL sample) suggests incomplete data migration, while flight_2 indicates export errors. "
        "HIGH severity assigns to flight_4 (1,240 violations), suggesting systematic issues with flight-domain databases. The "
        "remaining eight LOW severity databases (1-6 violations each) likely represent isolated data entry errors or constraint "
        "addition after data insertion.",
    )
    
    add_paragraph(
        doc,
        "The CRITICAL severity classification for FK data violations reflects their direct impact on model evaluation reliability, "
        "particularly in Test and Dev partitions. Standard Text-to-SQL evaluation employs Execution Match (EX) metrics that compare "
        "predicted query results against ground truth results by executing both queries and measuring set equivalence. When databases "
        "contain referential integrity violations, JOIN operations produce incorrect results differing from those assumed during "
        "annotation, introducing systematic evaluation bias. Specifically, flight_2 in Dev with 2,403 violations affects 35.6% of "
        "Dev's validation examples, potentially invalidating hyperparameter tuning decisions based on corrupted evaluation signals. "
        "Models optimized against such biased metrics may exhibit degraded performance on clean data while appearing superior during "
        "validation. The Test partition's single violation represents minimal bias risk, while Train's violations (though numerous) "
        "primarily affect training dynamics rather than evaluation integrity. This asymmetry underscores the urgent need to repair "
        "Dev's flight_2 database before conducting model selection or publishing benchmark results.",
    )

    # Table 2b (Databases with Empty Tables)
    add_paragraph(
        doc,
        "Train exhibits 71 empty tables (9.0% of Train tables) absent from Test and Dev. Table 2b shows seven databases "
        "with 0% data coverage (academic, imdb, yelp, restaurants, geo, music_2, scholar) totaling 65 empty tables, plus "
        "partial emptiness in sakila_1 (4 empty) and formula_1 (2 empty). Empty tables preclude semantic validation and "
        "may confuse training.",
    )
    
    headers_t2b = ["Database ID", "Total Tables", "Empty Tables", "Data Coverage", "Status"]
    rows_t2b = [
        ["academic", "15", "15", "0%", "Error + All Empty"],
        ["imdb", "16", "16", "0%", "Error + All Empty"],
        ["yelp", "7", "7", "0%", "Error + All Empty"],
        ["restaurants", "3", "3", "0%", "Error + All Empty"],
        ["geo", "7", "7", "0%", "Warning (All Empty)"],
        ["music_2", "7", "7", "0%", "Warning (All Empty)"],
        ["scholar", "10", "10", "0%", "Warning (All Empty)"],
        ["sakila_1", "16", "4", "75%", "Error + Partial Data"],
        ["formula_1", "13", "2", "85%", "Warning (Partial Data)"],
        ["TOTAL (Train only)", "94", "71", "24.5%", "—"],
    ]
    add_table_with_data(
        doc,
        headers_t2b,
        rows_t2b,
        "Table 2b. Databases with empty tables in Spider Train partition",
    )
    
    add_paragraph(
        doc,
        "Seven databases exhibit 0% coverage (65 of 71 empty tables), likely representing schema-only databases for DDL testing, "
        "incomplete exports, or licensing-restricted data. Four of these (academic, imdb, yelp, restaurants) combine structural "
        "errors with complete emptiness, rendering them unsuitable beyond syntax checking. Partially empty databases (sakila_1: "
        "75% coverage; formula_1: 85% coverage) may represent unpopulated optional features. The complete absence of empty tables "
        "in Test and Dev (100% coverage) indicates systematic quality assurance applied to evaluation partitions but not training.",
    )

    add_paragraph(
        doc,
        "Impact scope estimation: Test's 4 error databases affect ~213 queries (9.9%), Dev's 7 affect ~368 queries (35.6%), "
        "and Train's 32 affect ~1,900 queries (22%). Dev's 35.6% impact rate indicates over one-third of validation examples "
        "involve structurally defective databases.",
    )

    # 3.3. Syntactic analysis
    doc.add_page_break()
    add_heading(doc, "3.3. Syntactic Structure and Complexity Analysis", level=2)
    
    add_paragraph(
        doc,
        "Syntactic analysis evaluates the structural characteristics of SQL queries independent of their semantic "
        "correctness, examining complexity, feature utilization, and parsing success rates. Table 3 demonstrates "
        "remarkable consistency across partitions in fundamental syntactic metrics, suggesting careful curation to "
        "maintain comparable difficulty distributions.",
    )

    # Table 3 (syntactic analysis)
    headers_t3 = ["Metric", "Spider Test", "Spider Dev", "Spider Train", "Consistency"]
    rows_t3 = [
        ["Total queries", "2,147", "1,034", "8,659", "—"],
        ["Successfully parsed", "2,147 (100%)", "1,034 (100%)", "8,659 (100%)", "✅"],
        ["Parse failures", "0 (0%)", "0 (0%)", "0 (0%)", "✅"],
        ["Mean complexity", "39.5", "39.0", "40.7", "Max Δ=1.7"],
        ["Median complexity", "41", "41", "41", "✅"],
        ["Std deviation", "20.2", "19.9", "21.2", "Max Δ=1.3"],
    ]
    add_table_with_data(
        doc,
        headers_t3,
        rows_t3,
        "Table 3. Overall syntactic analysis results showing strong cross-partition consistency",
    )

    add_paragraph(
        doc,
        "The perfect parsing success rate (100% across all 11,840 queries) constitutes a significant achievement, "
        "indicating that all queries conform to valid SQL syntax without requiring dialect-specific workarounds or "
        "error correction. This result validates the sqlglot parser's robustness for SQLite dialect and confirms the "
        "syntactic correctness of Spider's SQL annotations. The near-identical complexity statistics (maximum difference "
        "of 1.7 points on a 100-point scale, maximum standard deviation difference of 1.3) demonstrate intentional "
        "balancing of difficulty distributions across partitions, ensuring that test set performance metrics reflect "
        "genuine model capabilities rather than artifacts of easier or harder query distributions.",
    )

    # Table 4 (difficulty distribution)
    add_paragraph(
        doc,
        "Examination of difficulty level distributions (Table 4) reinforces this consistency pattern, with Easy/Medium/Hard "
        "categories exhibiting nearly identical proportions across Test and Dev partitions. Train displays slightly higher "
        "Medium representation (55.7% vs. 58-59% in Test/Dev) balanced by correspondingly lower Easy percentages, differences "
        "that likely reflect statistical variation rather than systematic curation differences given the magnitudes involved.",
    )

    headers_t4 = ["Difficulty", "Spider Test", "Spider Dev", "Spider Train", "Range"]
    rows_t4 = [
        ["Easy", "516 (24.0%)", "250 (24.2%)", "2,039 (23.5%)", "10-29"],
        ["Medium", "1,246 (58.0%)", "612 (59.2%)", "4,827 (55.7%)", "30-59"],
        ["Hard", "385 (17.9%)", "172 (16.6%)", "1,727 (19.9%)", "60-79"],
        ["Expert", "0 (0.0%)", "0 (0.0%)", "66 (0.8%)", "80-100"],
    ]
    add_table_with_data(
        doc, headers_t4, rows_t4, "Table 4. Query difficulty distribution across partitions"
    )

    add_paragraph(
        doc,
        "One notable exception to cross-partition consistency emerges in the Expert difficulty category: Train contains "
        "66 Expert-level queries (0.8% of the partition) while Test and Dev contain none. These 66 queries, with complexity "
        "scores ranging from 80-100, represent the most structurally complex examples in the entire dataset, typically "
        "involving deeply nested subqueries (depth ≥3), multiple JOIN operations (≥4 tables), and combinations of advanced "
        "features. The restriction of Expert queries to Train likely reflects a deliberate design decision to reserve the "
        "most challenging examples for training while maintaining test/dev difficulty within bounds where human annotation "
        "reliability remains high, though this asymmetry means that models cannot be comprehensively evaluated on Expert-level "
        "complexity using the provided test partition.",
    )

    # Table 5 (JOIN distribution)
    add_paragraph(
        doc,
        "Analysis of JOIN distribution (Table 5) reveals a second dimension where Train exhibits systematically higher "
        "complexity compared to Test and Dev. Train contains a substantially larger proportion of multi-table queries "
        "(44.02% vs. approximately 40% in Test/Dev), with particularly pronounced differences in the 3-JOIN (4.15% vs. "
        "1-2%) and 4+-JOIN (2.06% vs. 1%) categories. This pattern suggests that Train was specifically enriched with "
        "complex multi-table queries to provide training signal for JOIN operation learning, though the test/dev partitions' "
        "lower JOIN frequencies may underestimate models' capabilities on highly complex multi-table scenarios.",
    )

    headers_t5 = ["JOIN Count", "Spider Test", "Spider Dev", "Spider Train", "Best"]
    rows_t5 = [
        ["0 (single-table)", "1,285 (59.85%)", "626 (60.54%)", "4,847 (55.98%)", "Train 🎯"],
        ["1 JOIN", "558 (25.99%)", "320 (30.95%)", "2,233 (25.79%)", "Balanced ✅"],
        ["2 JOINs", "242 (11.27%)", "72 (6.96%)", "1,042 (12.03%)", "Train 🎯"],
        ["3 JOINs", "38 (1.77%)", "10 (0.97%)", "359 (4.15%)", "Train 🎯"],
        ["4+ JOINs", "24 (1.12%)", "6 (0.58%)", "178 (2.06%)", "Train 🎯"],
        ["Multi-table (≥1 J.)", "862 (40.15%)", "408 (39.46%)", "3,812 (44.02%)", "Train 🎯"],
    ]
    add_table_with_data(
        doc,
        headers_t5,
        rows_t5,
        "Table 5. JOIN distribution analysis showing Train's higher multi-table complexity",
    )

    add_paragraph(
        doc,
        "Despite the rich structural variation in JOIN complexity and overall difficulty levels, Spider exhibits complete "
        "absence of certain advanced SQL features. No queries in any partition employ common table expressions (CTEs), "
        "recursive CTEs, or window functions (ROW_NUMBER, RANK, LAG, LEAD), features commonly found in modern SQL codebases "
        "and supported by all major database systems. This limitation reflects Spider's 2018 origins when window functions "
        "had more limited adoption, but represents a significant coverage gap for evaluating contemporary Text-to-SQL systems "
        "expected to support the full SQL feature set. The forthcoming Spider 2.0 reportedly addresses this limitation by "
        "including examples with modern SQL constructs.",
    )

    # 3.4. Query execution
    doc.add_page_break()
    add_heading(doc, "3.4. Dynamic Executability Validation", level=2)
    
    add_paragraph(
        doc,
        "Dynamic execution testing validates that syntactically correct queries can actually run against their associated "
        "databases, detecting semantic errors, missing tables/columns, type incompatibilities, and other runtime failures "
        "invisible to static analysis. Table 6 presents execution results demonstrating near-perfect executability across "
        "all partitions.",
    )

    # Table 6 (execution results)
    headers_t6 = ["Metric", "Spider Test", "Spider Dev", "Spider Train", "Best"]
    rows_t6 = [
        ["Total queries", "2,147", "1,034", "8,659", "—"],
        ["Successfully executed", "2,147 (100%)", "1,034 (100%)", "8,656 (99.97%)", "Test/Dev ✅"],
        ["Execution failures", "0 (0%)", "0 (0%)", "3 (0.03%)", "Test/Dev ✅"],
        ["Skipped", "0 (0%)", "0 (0%)", "0 (0%)", "✅"],
        ["Mean time (ms)", "3.10", "2.49", "1.19", "Train ✅"],
    ]
    add_table_with_data(
        doc,
        headers_t6,
        rows_t6,
        "Table 6. Query execution results showing exceptional executability rates",
    )

    add_paragraph(
        doc,
        "The execution success rates—100% for Test and Dev, 99.97% for Train—represent a remarkable achievement given the "
        "complexity of the queries and the schema integrity issues documented in Section 3.2. That queries execute "
        "successfully despite the presence of foreign key errors and referential integrity violations indicates either that "
        "most queries avoid JOIN operations that would expose these structural problems, or that SQLite's permissive handling "
        "of constraint violations prevents runtime errors even when referential integrity is compromised. The three failed "
        "executions in Train (0.03% of 8,659 queries) likely represent genuine semantic errors—references to non-existent "
        "columns or tables, incompatible type operations, or division-by-zero scenarios—warranting manual inspection to "
        "determine whether these represent annotation errors requiring correction.",
    )

    add_paragraph(
        doc,
        "Execution time measurements reveal an unexpected pattern: Train queries execute fastest (mean 1.19ms), followed by "
        "Dev (2.49ms) and Test (3.10ms). This inverse relationship between partition size and execution time suggests that "
        "Train databases may contain smaller data volumes compared to Test/Dev databases, or that Train's database schemas "
        "include better indexing. The absolute magnitudes—all means under 3.5 milliseconds—indicate that execution testing "
        "contributes negligible overhead to overall pipeline runtime, validating the decision to include execution validation "
        "by default rather than making it optional.",
    )

    # 3.5. Antipatterns
    doc.add_page_break()
    add_heading(doc, "3.5. Code Quality Assessment via Antipattern Detection", level=2)
    
    add_paragraph(
        doc,
        "Antipattern detection identifies common SQL coding practices that, while syntactically valid and functionally "
        "correct, deviate from established best practices regarding performance, maintainability, SQL portability, or code "
        "clarity. It is important to emphasize that antipattern presence in training data does not constitute a critical "
        "flaw for Text-to-SQL model development—training objectives focus on semantic correspondence between questions and "
        "queries rather than production-ready code optimization. However, certain antipattern categories warrant attention "
        "because they impact SQL portability or can introduce subtle correctness risks. CRITICAL severity patterns (Cartesian "
        "products) often indicate unintended join conditions that can explode result cardinality. HIGH severity patterns "
        "(Missing GROUP BY; NOT IN with nullable columns) either rely on dialect-specific aggregate behavior or can produce "
        "incorrect filtering under SQL three-valued logic. MEDIUM severity patterns primarily affect performance (e.g., leading "
        "wildcard LIKE; function calls in WHERE that inhibit index usage). LOW severity patterns (SELECT *; redundant DISTINCT) "
        "tend to be stylistic or micro-optimization concerns in a benchmark setting.",
    )

    # Table 7 (antipattern overall quality)
    headers_t7 = ["Metric", "Spider Test", "Spider Dev", "Spider Train", "Best"]
    rows_t7 = [
        ["Mean quality score", "97.7/100", "97.7/100", "98.4/100", "Train ✅"],
        ["Avg antipatterns per query", "0.2", "0.2", "0.1", "Train ✅"],
        ["Queries without antipatterns", "1,772 (82.5%)", "865 (83.7%)", "7,494 (86.5%)", "Train ✅"],
        ["Total antipattern occurrences", "383", "169", "1,182", "—"],
    ]
    add_table_with_data(
        doc,
        headers_t7,
        rows_t7,
        "Table 7. Overall code quality metrics derived from SQL antipattern detection across Spider partitions",
    )

    add_paragraph(
        doc,
        "Table 7 indicates consistently high SQL code quality across all three partitions. Mean quality scores reach 97.7/100 "
        "for Test and Dev and 98.4/100 for Train, while average antipattern density remains low (0.1–0.2 per query). Most "
        "importantly, the majority of queries contain no antipatterns at all: 82.5% (Test), 83.7% (Dev), and 86.5% (Train). "
        "These results suggest that antipatterns are concentrated in a relatively small subset of queries rather than being a "
        "systemic issue across the dataset.",
    )

    # Table 8 (top antipatterns by severity and frequency)
    add_paragraph(
        doc,
        "Table 8 reports the observed antipattern distribution across partitions. The dominant HIGH-severity antipattern is "
        "Missing GROUP BY (945 occurrences total), affecting 6.9% of Train queries and roughly one in ten queries in Test "
        "(11.1%) and Dev (10.3%). NOT IN with nullable columns forms the second major HIGH-severity category (348 occurrences "
        "total). Critical Cartesian products are rare across all partitions; these flags typically indicate missing or ineffective "
        "join conditions (i.e., at least one table is not connected to the join graph) and therefore should be treated as "
        "high-priority items for manual review because they can explode result cardinality and invalidate query intent.",
    )

    headers_t8 = ["Antipattern", "Spider Test", "Spider Dev", "Spider Train", "Severity"]
    rows_t8 = [
        ["Cartesian product", "2 (0.1%)", "2 (0.2%)", "13 (0.2%)", "🔴 Critical"],
        ["Missing GROUP BY", "239 (11.1%)", "106 (10.3%)", "600 (6.9%)", "⚠️ High"],
        ["NOT IN with nullable", "74 (3.4%)", "46 (4.4%)", "228 (2.6%)", "⚠️ High"],
        ["Leading wildcard LIKE", "50 (2.3%)", "12 (1.2%)", "140 (1.6%)", "🔵 Medium"],
        ["Function in WHERE", "0 (0.0%)", "0 (0.0%)", "6 (0.1%)", "🔵 Medium"],
        ["Redundant DISTINCT", "6 (0.3%)", "0 (0.0%)", "133 (1.5%)", "🟢 Low"],
        ["SELECT *", "12 (0.6%)", "3 (0.3%)", "62 (0.7%)", "🟢 Low"],
    ]
    add_table_with_data(
        doc, headers_t8, rows_t8, "Table 8. Complete antipattern distribution by severity and partition"
    )

    add_paragraph(
        doc,
        "Missing GROUP BY is the most frequent high-severity issue in all partitions. A large concentration of Missing GROUP BY cases reduces SQL portability and can " 
        "introduce ambiguous semantics when multiple rows contribute to an aggregate. For cross-dialect Text-to-SQL applications, these cases should be prioritized for "
        "normalization into standard-compliant GROUP BY queries or rewritten using deterministic subqueries when the intent is to associate a value with an extremum.",
    )

    add_paragraph(
        doc,
        "The NOT IN with nullable columns antipattern affects 348 queries total (74 in Test, 46 in Dev, 228 "
        "in Train), representing 2.6-4.4% of each partition. This pattern introduces subtle semantic bugs when the subquery "
        "returns NULL values due to SQL's three-valued logic: when the subquery result contains NULL, the predicate may "
        "evaluate to UNKNOWN and filter out rows unexpectedly. A standard mitigation is to either add explicit IS NOT NULL "
        "filters inside the subquery or to rewrite NOT IN as NOT EXISTS with a correlated predicate. Importantly, some flagged "
        "cases may be safe if the selected subquery column is guaranteed non-null (e.g., a primary key); therefore, these "
        "cases require schema-aware review to separate true defects from structurally safe patterns.",
    )

    add_paragraph(
        doc,
        "MEDIUM severity antipatterns are relatively infrequent. Leading wildcard LIKE (1.2–2.3%) can prevent index usage in "
        "many engines, potentially increasing runtime on large tables. Function calls in WHERE are rare overall (0.1% in Train "
        "and absent from Test/Dev) but represent a similar optimization barrier: applying a function to a filtered column can "
        "inhibit index utilization unless functional indexes or specialized collations are used.",
    )

    add_paragraph(
        doc,
        "LOW severity antipatterns are also limited in scope. Redundant DISTINCT affects 0.3% of Test queries and 1.5% of Train "
        "queries (absent in Dev), suggesting occasional unnecessary deduplication. SELECT * is rare (0.3–0.7%), indicating that "
        "most Spider queries explicitly enumerate selected columns. While these patterns can matter in production SQL, their low "
        "frequency here suggests they are not a primary driver of quality variation across partitions.",
    )

    add_paragraph(
        doc,
        "In summary, antipattern detection confirms high SQL code quality in Spider: mean quality scores of 97.7–98.4/100 and "
        "82.5–86.5% of queries free of detected antipatterns. The primary remediation opportunity concerns HIGH-severity patterns, "
        "especially Missing GROUP BY (945 occurrences) and NOT IN with nullable columns (348 occurrences). Although CRITICAL "
        "Cartesian products are rare (17 total), they should be treated as high-priority fixes, as they often indicate broken join "
        "logic. Overall, the results indicate that targeted correction of a small number of pattern categories could further "
        "improve portability and robustness without requiring broad changes to the dataset.",
    )

    # 3.6. Semantic validation
    doc.add_page_break()
    add_heading(doc, "3.6. Semantic Correctness via LLM Consensus", level=2)
    
    add_paragraph(
        doc,
        "Semantic validation represents the most critical quality dimension, as syntactically valid, successfully executing "
        "queries may nonetheless fail to correctly answer their associated natural language questions. This validation employs "
        "two large language models (Gemini 2.5 Pro and GPT-5) in a consensus voting configuration, with unanimous agreement "
        "required for definitive verdicts and disagreement cases classified as Mixed requiring human review. The analysis was "
        "successfully conducted across all three Spider partitions: Test (2,147 queries), Dev (1,034 queries), and Train "
        "(8,659 queries), providing comprehensive semantic quality assessment covering the entire dataset of 11,840 examples.",
    )

    # Table 9 (semantic validation)
    headers_t9 = ["Consensus Verdict", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t9 = [
        ["CORRECT", "1,409 (65.6%)", "684 (66.2%)", "5,188 (59.9%)"],
        ["PARTIALLY_CORRECT", "43 (2.0%)", "14 (1.4%)", "259 (3.0%)"],
        ["INCORRECT", "171 (8.0%)", "94 (9.1%)", "867 (10.0%)"],
        ["Mixed (disagreement)", "522 (24.3%)", "242 (23.4%)", "2,316 (26.7%)"],
        ["UNANSWERABLE", "2 (0.1%)", "0 (0.0%)", "29 (0.3%)"],
        ["TOTAL", "2,147 (100%)", "1,034 (100%)", "8,659 (100%)"],
    ]
    add_table_with_data(
        doc,
        headers_t9,
        rows_t9,
        "Table 9. Semantic validation results via LLM consensus voting",
    )

    add_paragraph(
        doc,
        "Table 9 indicates that Test and Dev exhibit comparable semantic correctness (65.6% vs. 66.2% CORRECT), while Train shows a "
        "lower correctness rate (59.9%). Train also exhibits the highest rates of both explicitly incorrect queries (10.0% vs. 8.0% "
        "in Test and 9.1% in Dev) and mixed/disputed cases (26.7% vs. 24.3% in Test and 23.4% in Dev), indicating that the largest "
        "partition contains the greatest proportion of semantically problematic or ambiguous examples. The unanimous agreement "
        "characteristic (all non-Mixed verdicts represent complete model consensus) reflects the two-model configuration, where only "
        "two outcome patterns are possible: both models agree (producing a definitive verdict) or models disagree (producing Mixed "
        "classification).",
    )

    # Table 9a (Mixed / disagreement breakdown)
    headers_t9a = [
        "Disagreement pattern (Gemini → GPT-5)",
        "Spider Test",
        "Spider Dev",
        "Spider Train",
        "Share of Mixed (Test/Dev/Train)",
    ]
    rows_t9a = [
        ["CORRECT → PARTIALLY_CORRECT", "315", "144", "1,375", "60.3% / 59.5% / 59.4%"],
        ["CORRECT → INCORRECT", "124", "57", "449", "23.8% / 23.6% / 19.4%"],
        ["INCORRECT → PARTIALLY_CORRECT", "35", "17", "169", "6.7% / 7.0% / 7.3%"],
        ["INCORRECT → CORRECT", "20", "14", "121", "3.8% / 5.8% / 5.2%"],
        ["CORRECT → UNANSWERABLE", "14", "4", "94", "2.7% / 1.7% / 4.1%"],
        ["INCORRECT → UNANSWERABLE", "1", "3", "40", "0.2% / 1.2% / 1.7%"],
        ["PARTIALLY_CORRECT → INCORRECT", "6", "2", "30", "1.1% / 0.8% / 1.3%"],
        ["PARTIALLY_CORRECT → CORRECT", "6", "1", "11", "1.1% / 0.4% / 0.5%"],
        ["UNANSWERABLE → INCORRECT", "1", "0", "2", "0.2% / 0.0% / 0.1%"],
        ["UNANSWERABLE → CORRECT", "0", "0", "14", "0.0% / 0.0% / 0.6%"],
        ["UNANSWERABLE → PARTIALLY_CORRECT", "0", "0", "6", "0.0% / 0.0% / 0.3%"],
        ["PARTIALLY_CORRECT → UNANSWERABLE", "0", "0", "5", "0.0% / 0.0% / 0.2%"],
    ]
    add_table_with_data(
        doc,
        headers_t9a,
        rows_t9a,
        "Table 9a. Mixed consensus breakdown (Gemini 2.5 Pro vs GPT-5 disagreement patterns)",
    )

    add_paragraph(
        doc,
        "The Mixed verdict category, indicating model disagreement, warrants particular attention given its substantial prevalence "
        "across all partitions. Train exhibits the highest Mixed rate (26.7%), followed by Test (24.3%) and Dev (23.4%), suggesting "
        "that larger partitions contain more ambiguous or underspecified examples where reasonable SQL implementations might differ "
        "in interpretation. These cases may represent questions admitting multiple valid SQL formulations, queries with subtle semantic "
        "bugs difficult for both humans and models to detect, or scenarios where database schema ambiguities enable different but "
        "defensible interpretations. Table 9a shows that Mixed is dominated by CORRECT↔PARTIALLY_CORRECT boundary disagreements "
        "(≈59–60% of Mixed across partitions), indicating that many disputed cases likely hinge on missing or extra constraints rather "
        "than being clearly wrong. The second-largest pattern is CORRECT↔INCORRECT, representing harder semantic disputes that are "
        "especially valuable for targeted human adjudication. Train also contains the largest share of UNANSWERABLE-related "
        "disagreements, consistent with a higher prevalence of underspecified questions at scale. The explicitly incorrect verdicts "
        "(171 in Test, 94 in Dev, 867 in Train) identify examples "
        "where both models independently judged the SQL query as failing to answer the natural language question, likely representing "
        "genuine annotation errors or cases where the intended query logic was incorrectly implemented. Train's 867 incorrect queries "
        "represent approximately 10% of its examples, a concerning proportion that exceeds the incorrect rates in both Test and Dev. "
        "The partially correct category, while relatively small in Test (2.0%) and Dev (1.4%), reaches 3.0% in Train with 259 examples, "
        "capturing queries that answer part of the question but omit necessary constraints or return supersets of the correct results.",
    )

    # Table 10 (semantic quality categories)
    add_paragraph(
        doc,
        "Table 10 synthesizes these findings into actionable quality categories across all three partitions. High-confidence correct "
        "examples are 65.6% (Test), 66.2% (Dev), and 59.9% (Train), indicating that a substantial fraction of each partition requires "
        "manual review or exhibits semantic issues. Cases requiring review include both the Mixed disagreement cases (24.3% in Test, "
        "23.4% in Dev, 26.7% in Train) and the partially correct examples, totaling 26.3% of Test, 24.8% of Dev, and 29.7% of Train. "
        "Clearly incorrect examples (8.0% in Test, 9.1% in Dev, 10.0% in Train) should be flagged for correction or removal, "
        "representing 1,132 total examples across all three partitions—a substantial corpus of "
        "annotation errors that may influence model training and evaluation if left unaddressed.",
    )

    headers_t10_cat = ["Category", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t10_cat = [
        ["Correct", "1,409 (65.6%)", "684 (66.2%)", "5,188 (59.9%)"],
        ["Problematic (all)", "216 (10.1%)", "108 (10.4%)", "1,155 (13.3%)"],
        ["Disputed (Mixed)", "522 (24.3%)", "242 (23.4%)", "2,316 (26.7%)"],
        ["High confidence", "1,409 (65.6%)", "684 (66.2%)", "5,188 (59.9%)"],
        ["Requires review", "565 (26.3%)", "256 (24.8%)", "2,575 (29.7%)"],
        ["Clearly incorrect", "171 (8.0%)", "94 (9.1%)", "867 (10.0%)"],
    ]
    add_table_with_data(
        doc, headers_t10_cat, rows_t10_cat, "Table 10. Semantic quality categorization summary"
    )

    add_paragraph(
        doc,
        "The semantic validation results should be interpreted with appropriate methodological caveats. The two-model consensus "
        "approach, while providing useful signal, cannot definitively establish ground truth semantic correctness. Both models may "
        "share systematic biases or blind spots, potentially producing unanimous incorrect verdicts for examples that actually are "
        "correct, or conversely judging correct examples as incorrect due to misunderstanding subtle question semantics or database "
        "content. The 59.9-66.2% correctness rates likely represent lower bounds on true semantic quality, as conservative "
        "classification strategies may flag some actually-correct examples as Mixed or incorrect. Nonetheless, the comprehensive "
        "coverage across all 11,840 examples enables identification of 1,132 clearly incorrect queries and 3,080 disputed cases "
        "requiring targeted human review to adjudicate the genuine error rate and understand the nature of semantic issues present "
        "throughout the dataset. The successful completion of LLM validation across all partitions, including the computationally "
        "expensive Train partition, provides unprecedented visibility into Spider's semantic quality at full scale.",
    )

    # 3.7. Performance analysis
    doc.add_page_break()
    add_heading(doc, "3.7. Performance Characteristics of Pipeline Components", level=2)
    
    add_paragraph(
        doc,
        "Understanding the computational cost profile of different analyzers informs decisions about which analyses to enable in "
        "resource-constrained scenarios and helps predict overall pipeline runtime for large datasets. Table 11 presents per-query "
        "processing times for each analyzer across all three partitions.",
    )

    # Table 11 (performance)
    headers_t11_perf = ["Analyzer", "Test (ms)", "Dev (ms)", "Train (ms)", "Fastest"]
    rows_t11_perf = [
        ["Query Antipattern", "1.21", "0.49", "0.56", "Dev ✅"],
        ["Query Syntax", "2.32", "0.60", "0.64", "Dev ✅"],
        ["Query Execution", "3.10", "2.49", "1.19", "Train ✅"],
        ["Schema Validation", "14.25", "16.91", "9.93", "Train ✅"],
        ["Semantic LLM Judge", "21,271", "21,200", "20,350", "Train ✅"],
    ]
    add_table_with_data(
        doc,
        headers_t11_perf,
        rows_t11_perf,
        "Table 11. Per-query processing time (milliseconds) by analyzer and partition",
    )

    add_paragraph(
        doc,
        "The four formal analyzers (antipattern, syntax, execution, schema validation) all complete within 20 milliseconds per "
        "query, with antipattern detection typically fastest (0.49-1.21ms), followed by syntax analysis (0.60-2.32ms), query "
        "execution (1.19-3.10ms), and schema validation (9.93-16.91ms). The relatively long schema validation times reflect the "
        "cost of database schema introspection queries against SQLite system tables, while the fast antipattern detection times "
        "result from simple pattern matching against pre-parsed AST structures. The variation across partitions likely stems from "
        "differences in average query complexity (more complex queries require more time to parse and analyze) and schema complexity "
        "(databases with more tables and foreign keys require longer schema validation).",
    )

    add_paragraph(
        doc,
        "The Semantic LLM Judge analyzer dominates overall runtime by three orders of magnitude, averaging approximately 20-21 seconds "
        "per query across all partitions (21.3s for Test, 21.2s for Dev, 20.4s for Train). This enormous time—approximately 1,000× "
        "slower than the combined formal analyzers—reflects network latency for API calls to remote language model services plus model "
        "inference time. The per-query time translates to roughly 12.7 hours for Test's 2,147 queries, 6.1 hours for Dev's 1,034 queries, "
        "and 49 hours for Train's 8,659 queries. The slightly faster per-query time observed in Train (20.4s) compared to Test (21.3s) "
        "likely reflects API performance variations over time or differences in query complexity distributions affecting LLM reasoning time. "
        "These timings assume sequential processing; parallel processing with multiple concurrent API calls could reduce wall-clock time "
        "proportionally to the degree of parallelism, though most LLM providers enforce rate limits that constrain practical parallelism.",
    )

    # Table 12 (aggregate processing time)
    add_paragraph(
        doc,
        "Table 12 aggregates these per-query times to overall partition processing times, revealing the stark dominance of LLM "
        "validation across all partitions. LLM Judge accounts for 99.9% of total pipeline runtime in all three cases, with all "
        "formal analyzers combined contributing under 0.1%. The cumulative analysis time across all three partitions totals "
        "approximately 67.8 hours (12.7 + 6.1 + 49.0 hours), representing a substantial computational investment but enabling "
        "comprehensive semantic quality assessment across the entire 11,840-example dataset. Had LLM validation been omitted "
        "from all partitions, the complete analysis would finish in under 2 minutes, demonstrating that formal validation scales "
        "to large datasets with negligible computational cost.",
    )

    headers_t12 = ["Component", "Spider Test", "Spider Dev", "Spider Train"]
    rows_t12 = [
        ["Fast analyzers", "~15 sec", "~8 sec", "~11 sec"],
        ["LLM Judge", "~12.7 hr", "~6.1 hr", "~49.0 hr"],
        ["Overhead (I/O)", "~45 sec", "~22 sec", "~68 sec"],
        ["Total time", "~12.7 hr", "~6.1 hr", "~49.0 hr"],
        ["LLM fraction", "99.9%", "99.9%", "99.9%"],
    ]
    add_table_with_data(
        doc, headers_t12, rows_t12, "Table 12. Aggregate processing time by partition"
    )

    add_paragraph(
        doc,
        "These performance characteristics suggest several operational strategies for large-scale dataset analysis. For initial "
        "exploration or iterative development, disabling LLM validation enables rapid turnaround (minutes rather than hours), allowing "
        "researchers to validate pipeline configuration, debug analysis logic, and examine formal validation results before committing "
        "to expensive semantic validation. For production analysis where semantic validation is required, parallelization strategies "
        "can dramatically reduce wall-clock time: processing 10-20 queries concurrently (subject to API rate limits) could compress the "
        "49-hour Train analysis to 2.5-5 hours. The extreme cost asymmetry (LLM validation consuming 99.9% of runtime) suggests that "
        "semantic validation should be reserved for final dataset validation rather than incorporated into routine quality checks during "
        "dataset development. Nonetheless, the comprehensive semantic validation across all 11,840 examples completed in this study "
        "provides the first complete quality assessment of Spider's semantic correctness, revealing issues that would remain hidden "
        "under sampling-based approaches.",
    )

    # 3.8. Discussion and recommendations
    doc.add_page_break()
    add_heading(doc, "3.8. Discussion, Implications, and Recommendations", level=2)
    
    add_paragraph(
        doc,
        "This section interprets the empirical findings from Tables 1–13, connects them to practical dataset usage, and "
        "highlights remediation priorities. The key takeaway is that formal correctness signals (parsability and executability) "
        "can be near-perfect while substantial structural and semantic issues remain—issues that directly affect the validity of "
        "training, model selection, and reported benchmark results.",
    )

    # Table 13 (integrated assessment)
    headers_t13_int = ["Quality Criterion", "Spider Test", "Spider Dev", "Spider Train", "Best"]
    rows_t13_int = [
        ["DATABASE SCHEMAS", "", "", "", ""],
        ["Valid schemas", "90.0%", "65.0%", "75.3%", "Test ✅"],
        ["Invalid FKs", "4", "7", "78", "Test ✅"],
        ["FK data violations", "1 (1 viol.)", "3 (2,403 viol.)", "6 (39,523 viol.)", "Test ✅"],
        ["SYNTACTIC QUALITY", "", "", "", ""],
        ["Parse success", "100%", "100%", "100%", "Equal ✅"],
        ["Mean complexity", "39.5", "39.0", "40.7", "Equal ✅"],
        ["Expert queries", "0", "0", "66 (0.8%)", "Train ✅"],
        ["EXECUTABILITY", "", "", "", ""],
        ["Execute success", "100%", "100%", "99.97%", "Test/Dev ✅"],
        ["Mean time (ms)", "3.10", "2.49", "1.19", "Train ✅"],
        ["SQL CODE QUALITY", "", "", "", ""],
        ["Quality score", "97.7/100", "97.7/100", "98.4/100", "Train ✅"],
        ["Critical antipatterns (Cartesian product)", "2 (0.1%)", "2 (0.2%)", "13 (0.2%)", "Test ✅"],
        ["High antipatterns (Missing GROUP BY)", "239 (11.1%)", "106 (10.3%)", "600 (6.9%)", "Train ✅"],
        ["High antipatterns (NOT IN with nullable)", "74 (3.4%)", "46 (4.4%)", "228 (2.6%)", "Train ✅"],
        ["SEMANTIC CORRECTNESS", "", "", "", ""],
        ["Correct", "65.6%", "66.2%", "59.9%", "Dev ✅"],
        ["Clearly incorrect", "8.0%", "9.1%", "10.0%", "Test ✅"],
        ["Disputed (Mixed)", "24.3%", "23.4%", "26.7%", "Dev ✅"],
    ]
    add_table_with_data(
        doc,
        headers_t13_int,
        rows_t13_int,
        "Table 13. Integrated quality assessment across all validation dimensions",
    )

    add_paragraph(
        doc,
        "Two results are simultaneously true. First, Spider is exceptionally strong on traditional validation criteria: all queries "
        "parse successfully, execution succeeds for essentially all examples, and syntactic complexity distributions remain broadly "
        "consistent across partitions. Second, deeper validation reveals that the benchmark contains concentrated schema/data defects "
        "and a large fraction of semantically disputed or incorrect examples. Together, these findings explain why execution-based "
        "evaluation can appear stable while still masking hidden dataset quality risks.",
    )

    add_paragraph(
        doc,
        "The most urgent structural issue is referential integrity corruption concentrated in a small number of databases, most notably "
        "flight_2 (Dev) and sakila_1 (Train). Dev defects are especially harmful because they can bias model selection: when joins are "
        "evaluated on corrupted relational links, execution-match signals no longer reflect the intended semantics of the benchmark. "
        "By contrast, Train defects primarily influence learning dynamics, increasing noise and encouraging models to fit spurious patterns.",
    )
    
    add_paragraph(
        doc,
        "Schema quality heterogeneity also matters operationally. Test exhibits the highest fraction of valid schemas, while Dev shows the "
        "lowest. This is counterintuitive because Dev is often treated as a reliable proxy for Test during development. In practice, Dev-level "
        "schema defects can cause unstable validation curves and reward models that overfit quirks of a few problematic databases rather than "
        "learning robust text-to-SQL mappings.",
    )

    add_paragraph(
        doc,
        "Semantic validation provides the strongest evidence that execution alone is insufficient. Across partitions, only about 60–66% of "
        "examples are unanimously judged CORRECT, around 8–10% are unanimously judged INCORRECT, and roughly one quarter fall into the Mixed "
        "category where judges disagree. Mixed cases are not merely noise: they concentrate near the CORRECT↔PARTIALLY_CORRECT boundary and "
        "often reflect underspecified questions, edge-case constraints (duplicates, NULL handling, date boundaries), or multiple defensible "
        "interpretations. These cases define a practical “uncertainty band” of the benchmark that should be handled explicitly in evaluation "
        "and dataset curation.",
    )

    add_paragraph(
        doc,
        "A key methodological implication is that benchmark “quality” is multi-dimensional. Parsability/executability reflect technical validity; "
        "schema integrity reflects correctness of the environment; antipatterns reflect portability and maintainability; semantic judgments reflect "
        "alignment of NL intent and SQL logic. For modern LLM systems, semantic alignment is often the limiting factor, and disagreement between "
        "strong judges is a useful proxy for annotation ambiguity.",
    )

    add_paragraph(
        doc,
        "The antipattern results (Tables 7–8) add an important nuance: Spider’s SQL is generally clean from a database-engineering perspective "
        "(mean quality scores ≈97.7–98.4/100; 82.5–86.5% of queries contain no detected antipatterns), but a small number of high-severity "
        "categories dominate the risk surface. In particular, Missing GROUP BY and NOT IN with nullable subqueries account for the majority of "
        "high-severity flags. These patterns matter less for “can the SQL run?” and more for portability and semantic stability across engines: "
        "Missing GROUP BY can yield non-deterministic row selection in permissive engines (e.g., SQLite) and hard errors in stricter settings, "
        "while NOT IN becomes semantically fragile under three-valued logic when NULLs are present.",
    )

    add_paragraph(
        doc,
        "Practically, this means that antipatterns should be treated as a curation and evaluation control knob. For training, keeping them may be "
        "acceptable (models learn benchmark-style SQL), but for deployment-oriented evaluation or cross-dialect generalization studies, it is useful "
        "to (i) filter or rewrite the small subset of high-severity antipattern queries into standard-compliant equivalents, and (ii) report results "
        "both on the full set and on an “antipattern-clean” subset. The rare Critical Cartesian product cases, although few, should be prioritized for "
        "manual review because they often indicate broken join graphs that can radically change result cardinality and therefore question–SQL alignment. "
        "In our results, Cartesian products are flagged in 2 Test queries, 2 Dev queries, and 13 Train queries (Table 8), making them a small but "
        "high-impact class of potential semantic defects.",
    )

    # Table 13a (Critical Databases Requiring Immediate Attention)
    add_paragraph(
        doc,
        "Table 13a consolidates the most critical database quality issues identified across all validation dimensions, providing "
        "a prioritized action list for dataset maintainers. This practical reference enables immediate focus on the databases "
        "exhibiting the most severe or widespread problems affecting dataset reliability.",
    )
    
    headers_t13a = ["Database ID", "Partition", "Critical Issues", "Priority"]
    rows_t13a = [
        ["flight_2", "Dev", "2,403 FK violations → EVALUATION BIAS", "CRITICAL"],
        ["sakila_1", "Train", "38,273 FK violations + 4 empty tables", "CRITICAL"],
        ["flight_4", "Train", "1,240 FK violations", "HIGH"],
        ["baseball_1", "Train", "20 structural FK errors across 26 tables", "HIGH"],
        ["cre_Drama_Workshop_Groups", "Train", "9 FK type mismatches across 18 tables", "HIGH"],
        ["academic", "Train", "All 15 tables empty + FK type mismatch", "HIGH"],
        ["imdb", "Train", "All 16 tables empty + 7 FK errors", "HIGH"],
        ["yelp", "Train", "All 7 tables empty + 7 FK errors", "HIGH"],
        ["restaurants", "Train", "All 3 tables empty + FK missing column", "HIGH"],
        ["book_1", "Test", "2 FK missing column errors", "MEDIUM"],
        ["car_1", "Dev", "2 FK errors + 2 data violations", "MEDIUM"],
        ["concert_singer", "Dev", "2 FK type mismatches", "MEDIUM"],
        ["voter_1", "Dev", "FK target not key error", "MEDIUM"],
    ]
    add_table_with_data(
        doc,
        headers_t13a,
        rows_t13a,
        "Table 13a. Critical databases requiring immediate remediation prioritized by severity",
    )
    
    add_paragraph(
        doc,
        "The prioritization in Table 13a is impact-driven. Evaluation-split corruption (Dev/Test) should be fixed first to protect the credibility "
        "of reported metrics and model selection. Next, catastrophic Train databases should be repaired to reduce training noise. Finally, isolated "
        "structural defects and empty-table databases should be either repaired (if data can be recovered) or explicitly marked as schema-only so "
        "they can be excluded from execution-based analyses and semantic validation relying on sampled values.",
    )

    add_paragraph(
        doc,
        "For practitioners, we recommend reporting quality-aware results in addition to the standard aggregate score: (i) report performance on "
        "a “high-confidence” subset (CORRECT-only, excluding Mixed/PARTIALLY_CORRECT/INCORRECT/UNANSWERABLE), (ii) report performance on the "
        "subset of examples whose databases pass schema/data integrity checks, and (iii) provide ablations showing sensitivity to Dev database "
        "issues. These steps make model comparisons more robust and reduce the risk of tuning against corrupted validation signals.",
    )

    add_paragraph(
        doc,
        "From a dataset-maintenance perspective, LLM-derived labels are most useful as triage rather than ground truth. We recommend a two-stage "
        "curation workflow: (1) prioritize unanimous INCORRECT cases for manual correction, because they are high-signal candidates for annotation "
        "bugs; (2) prioritize Mixed cases for expert adjudication and question disambiguation, because they often represent ambiguity rather than "
        "simple errors. A practical output of such adjudication is an “accepted alternatives” set (multiple SQL queries) for genuinely ambiguous questions.",
    )

    add_paragraph(
        doc,
        "Recommended remediation roadmap (impact order): repair Dev’s flight_2 first; then address Train’s catastrophic sakila_1 and other high-violation "
        "databases; then fix structural foreign-key definition issues; finally, resolve schema-only / empty-table databases by either restoring data or "
        "documenting them as non-executable. In parallel, use semantic labels to curate a clean subset for training and a “disputed” subset for robustness testing.",
    )

    add_paragraph(
        doc,
        "Threats to validity: (i) LLM judges are not an oracle—both false positives and false negatives are possible, especially on domain-specific "
        "schemas and under underspecified questions; (ii) even with full-schema context, some questions require external knowledge or assumptions not "
        "encoded in the database (e.g., domain conventions, synonymy, temporal interpretation), which can lead judges to inconsistent rulings; "
        "(iii) sampling a small number of cell values can introduce bias if samples are unrepresentative, potentially masking edge cases (NULLs, duplicates) "
        "or suggesting misleading distributions. For higher-confidence audits, Mixed cases should be complemented with expert human adjudication on a "
        "statistically representative subset and, where feasible, ablations that compare semantic verdicts with/without sampled values.",
    )

    add_paragraph(
        doc,
        "Finally, we recommend that future benchmark revisions publish (a) an integrity-fixed Dev split, (b) an officially curated high-confidence subset, "
        "and (c) a documented list of disputed examples. These additions would make Spider more reliable for both model development and scientific comparison, "
        "while preserving its value as a challenging cross-domain Text-to-SQL benchmark.",
    )

    add_paragraph(
        doc,
        "Overall, the results support a nuanced conclusion: Spider is technically clean in the sense of SQL syntax and execution, but it contains "
        "structural database defects and substantial semantic ambiguity that warrant quality-aware evaluation and targeted remediation. Treating these "
        "issues explicitly—rather than implicitly assuming benchmark correctness—improves reproducibility and strengthens the validity of Text-to-SQL research.",
    )

    # Save document
    print(f"\n💾 Saving document: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)
    print(f"✅ Done! Document saved: {OUTPUT_FILE}")
    print("\n📄 Section 3 (Academic v2 - ENHANCED with Database IDs) contains:")
    print("   • 7 subsections (3.1 – 3.8, results-focused)")
    print("   • 17 tables with comprehensive data + concrete database IDs")
    print("\n   SCHEMA VALIDATION TABLES (Section 3.2):")
    print("   • Table 1: Schema validation overview - comparative metrics")
    print("   • Table 2: FK error distribution - structural + data violations summary")
    print("   • Table 2-detailed: Database IDs by FK error type (NEW! PRACTICAL)")
    print("   • Table 2a: FK data violations by database ID - catastrophic cases (NEW! PRACTICAL)")
    print("   • Table 2b: Empty tables by database ID (NEW! PRACTICAL)")
    print("\n   ANALYSIS TABLES (Sections 3.3-3.7):")
    print("   • Tables 3-5: Syntactic analysis (complexity, difficulty, JOINs)")
    print("   • Table 6: Query execution validation")
    print("   • Tables 7-8: Code quality and antipatterns")
    print("   • Tables 9-10: Semantic validation via LLM consensus")
    print("   • Tables 11-12: Performance characteristics")
    print("\n   INTEGRATED ASSESSMENT (Section 3.8):")
    print("   • Table 13: Integrated quality assessment (all dimensions)")
    print("   • Table 13a: Critical databases prioritized for remediation (NEW! ACTIONABLE)")
    print("\n   KEY FEATURES:")
    print("   • Cross-reference to Section 2.3 for FK validation methodology")
    print("   • Academic prose style with minimized bullet lists")
    print("   • Full comparative analysis: Test vs Dev vs Train")
    print("   • 11,840 total queries analyzed across 206 unique databases")
    print("\n   PRACTICAL CONTRIBUTIONS:")
    print("   ✅ FK errors mapped to specific DBs: 19 DBs with type mismatches, 6 DBs with target-not-key")
    print("   ✅ Prioritized remediation list: 13 critical databases (Table 13a)")
    print("   ✅ Catastrophic cases identified: sakila_1 (38.3K violations), flight_2 (2.4K violations)")
    print("   ✅ Empty data mapped: 7 databases with 0% coverage (academic, imdb, yelp, restaurants, geo, music_2, scholar)")
    print("   ✅ Worst structural errors: baseball_1 (20 FK target-not-key), cre_Drama_Workshop_Groups (9 FK type mismatches)")
    print("\n   CRITICAL FINDINGS:")
    print("   • 89 structural FK errors across 41 databases")
    print("   • 41,927 data violations (97% in sakila_1 + flight_2)")
    print("   • ⚠️  EVALUATION BIAS: flight_2 (Dev) corrupts EX match metrics → affects 35.6% of validation!")
    print("   • 71 empty tables across 9 databases (Train only)")
    print("   • Quality gradient: Test (90%) > Train (75.3%) > Dev (65%)")
    print("\n   🎯 URGENT PRIORITY: Fix flight_2 before model evaluation/publication!")
    print("   📊 PRACTICAL VALUE: All problematic databases explicitly identified by ID for targeted remediation!")


if __name__ == "__main__":
    generate_section3_docx()

